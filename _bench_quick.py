#!/usr/bin/env python3
import os, sys, time, subprocess

pdf = sys.argv[1]
print("PDF:", pdf)

# pdf2docx only
out = "/tmp/_bench_pdf2docx.docx"
t0 = time.time()
from pdf2docx import Converter
cv = Converter(pdf)
cv.convert(out, start=0, end=None)
cv.close()
print(f"pdf2docx: {time.time()-t0:.1f}s size={os.path.getsize(out):,}")

from docx import Document
d = Document(out)
imgs = sum(1 for r in d.part.rels.values() if "image" in r.reltype)
chars = sum(len(p.text) for p in d.paragraphs)
print(f"images={imgs} text_chars={chars} tables={len(d.tables)} sections={len(d.sections)}")

# pdftoppm only
t0 = time.time()
subprocess.run(["pdftoppm", "-png", "-r", "150", pdf, "/tmp/bench_page"], check=True, capture_output=True, timeout=300)
n = len([f for f in os.listdir("/tmp") if f.startswith("bench_page") and f.endswith(".png")])
print(f"pdftoppm 150dpi: {time.time()-t0:.1f}s pages={n}")

# OCR 1 page
import pytesseract
from PIL import Image
pages = sorted(f for f in os.listdir("/tmp") if f.startswith("bench_page") and f.endswith(".png"))
t0 = time.time()
t = pytesseract.image_to_string(Image.open(os.path.join("/tmp", pages[0])), lang="chi_sim", config="--psm 6 --oem 1")
print(f"OCR page1: {time.time()-t0:.1f}s chars={len(t.strip())}")
print(f"OCR 52pg estimate: {(time.time()-t0)*52:.0f}s")