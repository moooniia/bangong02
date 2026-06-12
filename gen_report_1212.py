#!/usr/bin/env python3
"""生成 1212 项目验收报告 Word 文档。"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

OUT = r"C:\Users\paz\Desktop\1212验收报告.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(11)
    run.bold = bold
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = str(val)
            for p in table.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.1)
    sec.right_margin = Inches(1.1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("1212.pdf 扫描件转 Word 调试验收报告")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "黑体"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("对比参照：WPS 转换版（1212.docx）｜生成日期：2026年6月11日")
    sr.font.size = Pt(10)
    sr.font.name = "宋体"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    doc.add_paragraph()

    add_heading(doc, "一、验收结论", 1)
    add_para(doc, "调试阶段目标已基本达成。与 WPS 版相比，我们在「可编辑文字覆盖率、无整页截图、封面版式、水印过滤」上达到可上线水准；剩余约 6% 文字差距主要为 WPS 保留的 OCR 乱码/水印碎片，继续追赶性价比低，建议暂停字数优化，转入上线准备。")

    add_heading(doc, "二、样张与文件", 2)
    add_table(doc, ["文件", "说明", "路径"], [
        ["原稿 PDF", "12 页验收样张，涵盖封面/正文/签章/表格等", r"C:\Users\paz\Desktop\1212.pdf"],
        ["WPS 参照", "WPS 自带 PDF 转 Word", r"C:\Users\paz\Desktop\1212.docx"],
        ["我们输出", "火山 OCR + 自研排版（最新版）", r"C:\Users\paz\Desktop\1212_full.docx"],
        ["参考合同", "版式参照样例", r"Desktop\2024中国移动…服务合同 29.74254.docx"],
    ])

    add_heading(doc, "三、核心数据对比", 1)
    add_table(doc, ["指标", "WPS", "我们", "说明"], [
        ["可编辑文字量", "8357 字", "7843 字（93.8%）", "接近 WPS，缺口多为乱码非正文"],
        ["整页截图", "0 张", "0 张", "原 14 张，已消除"],
        ["Word 表格标签", "32", "18", "WPS 含嵌套表，计数偏高；表意内容已接近"],
        ["嵌入图片", "39", "18", "我们仅保留章/二维码/红章，更精简"],
        ["文件体积", "1.04 MB", "0.59 MB", "无大图截图后更小"],
        ["签章/第八条等关键词", "全有", "全有", "签字页、签章页、大表页均已覆盖"],
    ])

    add_heading(doc, "四、第 1 页（封面）专项", 2)
    add_table(doc, ["检查项", "结果"], [
        ["蓝色盖章编号 #008FEF", "通过"],
        ["合同编号 CMCCTD-SS-202400146", "通过"],
        ["黑体大字「服务合同」", "通过"],
        ["甲乙方【】格式", "通过"],
        ["章图 + 二维码（2 个小图）", "通过"],
        ["水印文字滤除", "通过"],
        ["标题居中、版式优于 WPS", "通过"],
    ])
    add_para(doc, "第 1 页自验：0 项问题。")

    add_heading(doc, "五、本轮主要改动", 1)
    items = [
        "签章页策略调整：由「整页截图」改为「可编辑文字 + 红章透明抠图」。",
        "表格策略调整：去掉 0.78 质量门槛，差表也尽量输出原生 Word 表格。",
        "截图回退收紧：仅 OCR 几乎为空时才整页截图。",
        "封面小图：火山 CDN 过期时从 PDF 坐标裁切补章图/二维码。",
        "部署冒烟：默认不再跑 pdf→word，避免每次 deploy 多烧 Volc 页数。",
        "OCR 缓存：1 小时内复用明细，避免同 PDF 重复扣费。",
    ]
    for i, t in enumerate(items, 1):
        add_para(doc, f"{i}. {t}")

    add_heading(doc, "六、火山 OCR 费用说明", 1)
    add_para(doc, "计费单位是「页」而非「次」。1212.pdf 完整 OCR 一次 = 12 页 = 0.24 元（按量 0.02 元/页）。免费 500 页在一天内耗尽，主因是调试时同一 PDF 被多次 OCR（转换 + 拉明细 + 诊断 + 每次部署冒烟等），非密钥泄露。")
    add_table(doc, ["场景", "页数", "费用"], [
        ["转 1 次 1212.pdf", "12 页", "约 0.24 元"],
        ["调试 10 轮（含明细+转换）", "约 240 页", "约 4.8 元"],
        ["继续抠 6% 字数（预估）", "数十～上百页试错", "收益极低"],
    ])

    add_heading(doc, "七、性价比判断与建议", 1)
    add_para(doc, "不建议继续追求 93.8% → 98% 文字量。代价：反复 OCR 费用、过滤规则变复杂、水印可能回流；收益：用户几乎感知不到。", bold=True)
    add_para(doc, "建议优先做（性价比高）：")
    for t in [
        "真实用户路径验收：上传 → 转换 → 下载（一次约 0.24 元）。",
        "OCR 失败/额度耗尽时前端明确提示，避免静默掉 Tesseract。",
        "改代码用缓存明细，少打重复 OCR。",
    ]:
        add_para(doc, f"• {t}")
    add_para(doc, "建议暂缓（性价比低）：")
    for t in [
        "为追字数放宽噪声过滤（易倒退）。",
        "全面追赶 WPS 嵌套表计数。",
        "逐字抠 OCR 错字到 100% 一致。",
    ]:
        add_para(doc, f"• {t}")

    add_heading(doc, "八、当前状态", 1)
    add_table(doc, ["项目", "状态"], [
        ["代码", "已部署至 139.196.28.78"],
        ["火山 OCR", "计费已开通，实测可用"],
        ["输出样例", r"Desktop\1212_full.docx"],
        ["调试阶段", "可收尾，建议转上线准备"],
    ])

    add_para(doc, "")
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = foot.add_run("办公工具箱 bangong02.com｜火山 OCR + volc_ocr.py")
    fr.font.size = Pt(9)
    fr.font.name = "宋体"
    fr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    main()