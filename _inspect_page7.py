#!/usr/bin/env python3
from docx import Document
import zipfile

path = r"C:\Users\paz\Desktop\P T W 测试\page_7_p3.docx"
doc = Document(path)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    print(f"{i:02d} | {t[:95]}")
print("tables", len(doc.tables))
with zipfile.ZipFile(path) as z:
    xml = z.read("word/document.xml").decode("utf-8")
    print("anchors", xml.count("wp:anchor"))
    print("page_breaks", xml.count('w:type="page"'))