#!/usr/bin/env python3
import zipfile
from docx import Document

path = r"C:\Users\paz\Desktop\P T W 测试\page_6_v0.10.6.docx"
with zipfile.ZipFile(path) as z:
    xml = z.read("word/document.xml").decode("utf-8")
print("page_breaks", xml.count('w:type="page"'))
print("sectPr", xml.count("w:sectPr"))
print("lastRenderedPageBreak", xml.count("w:lastRenderedPageBreak"))

doc = Document(path)
print("paragraphs", len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs):
    print(f"p{i}: {p.text[:80]!r}")

print("tables", len(doc.tables))
if doc.tables:
    t = doc.tables[0]
    print("rows", len(t.rows), "cols", len(t.columns))
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for ri in [0, 1, 27, 53]:
        if ri < len(t.rows):
            tr = t.rows[ri]._tr
            heights = tr.findall(f".//{ns}trHeight")
            hval = heights[0].get(f"{ns}val") if heights else None
            print(f"row{ri} height={hval}")
            # cell text len
            texts = [c.text[:30] for c in t.rows[ri].cells[:3]]
            print(f"  cells: {texts}")