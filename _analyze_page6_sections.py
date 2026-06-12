#!/usr/bin/env python3
import zipfile
import re
from docx import Document

path = r"C:\Users\paz\Desktop\P T W 测试\page_6_out.docx"
with zipfile.ZipFile(path) as z:
    xml = z.read("word/document.xml").decode("utf-8")

# section orientations
parts = xml.split("<w:sectPr")
print("sections", len(parts) - 1)
for i, part in enumerate(parts[1:], 1):
    orient = "landscape" if "w:orient w:val=\"landscape\"" in part else "portrait"
    pg = re.search(r'w:pgSz[^>]*w:w="(\d+)"[^>]*w:h="(\d+)"', part)
    dims = f"{pg.group(1)}x{pg.group(2)}" if pg else "?"
    print(f" sect{i}: {orient} pgSz={dims}")

doc = Document(path)
print("doc.sections", len(doc.sections))
for i, s in enumerate(doc.sections):
    from docx.enum.section import WD_ORIENT
    o = "landscape" if s.orientation == WD_ORIENT.LANDSCAPE else "portrait"
    print(f" sec{i}: {o} w={s.page_width} h={s.page_height} margins T={s.top_margin} B={s.bottom_margin}")

if doc.tables:
    t = doc.tables[0]
    lens = []
    for row in t.rows:
        max_len = max(len(c.text) for c in row.cells)
        lens.append(max_len)
    print("max cell lens", max(lens), "avg", sum(lens)/len(lens))
    long_rows = sum(1 for l in lens if l > 40)
    print("rows with cell>40chars", long_rows)