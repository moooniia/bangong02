from docx import Document
from docx.enum.section import WD_ORIENT

def inspect(path, label):
    doc = Document(path)
    sec = doc.sections[0]
    orient = "Portrait" if sec.orientation == WD_ORIENT.PORTRAIT else "Landscape"
    print(f"=== {label} ===")
    print(f"  Page: {round(sec.page_width.inches,2)}in x {round(sec.page_height.inches,2)}in ({orient})")
    print(f"  Margins T/B/L/R: {round(sec.top_margin.inches,2)} / {round(sec.bottom_margin.inches,2)} / {round(sec.left_margin.inches,2)} / {round(sec.right_margin.inches,2)}")
    print(f"  Sections: {len(doc.sections)}, Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
    for i, t in enumerate(doc.tables):
        rows = len(t.rows)
        cols = len(t.columns)
        cells = []
        for row in t.rows[:2]:
            for cell in row.cells[:4]:
                txt = cell.text.strip()[:20]
                if txt:
                    cells.append(txt)
        print(f"  Table[{i}]: {rows}r x {cols}c  head: {' | '.join(cells[:6])}")
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    print(f"  Non-empty paragraphs: {len(texts)}")
    for t in texts[:20]:
        print(f"    P: {t[:90]}")

inspect(r"C:\Users\paz\Desktop\P T W 测试\page_4_v0.10.13.docx", "v0.10.13 (ours)")
print()
inspect(r"C:\Users\paz\Desktop\P T W 测试\page_4.docx", "WPS standard")
