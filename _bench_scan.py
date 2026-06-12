#!/usr/bin/env python3
"""Benchmark scanned PDF conversion paths on server."""
import os
import subprocess
import sys
import tempfile
import time

pdf = sys.argv[1] if len(sys.argv) > 1 else "/home/toolbox/fixtures/2.pdf"

print("PDF:", pdf, "size:", os.path.getsize(pdf))
print("pages:", end=" ")
subprocess.run(["pdfinfo", pdf], capture_output=True, text=True)
r = subprocess.run(["pdfinfo", pdf], capture_output=True, text=True)
for line in r.stdout.splitlines():
    if "Pages:" in line:
        print(line.strip())

# 1) pdf2docx embed images
out = "/tmp/_bench_pdf2docx.docx"
t0 = time.time()
from pdf2docx import Converter

cv = Converter(pdf)
cv.convert(out, start=0, end=None)
cv.close()
t_pdf2docx = time.time() - t0
print(f"pdf2docx: {t_pdf2docx:.1f}s, size={os.path.getsize(out):,}")

from docx import Document

d = Document(out)
imgs = sum(1 for rel in d.part.rels.values() if "image" in rel.reltype)
chars = sum(len(p.text) for p in d.paragraphs)
print(f"  images={imgs}, text_chars={chars}, tables={len(d.tables)}")

# 2) pdftoppm
tmp = tempfile.mkdtemp(prefix="bench_")
t0 = time.time()
subprocess.run(
    ["pdftoppm", "-png", "-r", "150", pdf, os.path.join(tmp, "page")],
    check=True,
    capture_output=True,
    timeout=300,
)
pages = sorted(f for f in os.listdir(tmp) if f.endswith(".png"))
t_ppm = time.time() - t0
print(f"pdftoppm 150dpi: {t_ppm:.1f}s, pages={len(pages)}")

# 3) OCR per page
import pytesseract
from PIL import Image

t0 = time.time()
total_chars = 0
for i, name in enumerate(pages):
    t = pytesseract.image_to_string(
        Image.open(os.path.join(tmp, name)),
        lang="chi_sim",
        config="--psm 6 --oem 1",
    )
    total_chars += len(t.strip())
    if i == 0:
        t_first = time.time() - t0
t_ocr_all = time.time() - t0
print(f"OCR all {len(pages)} pages: {t_ocr_all:.1f}s (first page {t_first:.1f}s)")
print(f"  total_chars={total_chars}")
est_52 = t_ppm / max(len(pages), 1) * 52 + t_ocr_all / max(len(pages), 1) * 52
print(f"estimate 52-page contract OCR path: ~{est_52:.0f}s")