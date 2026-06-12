#!/usr/bin/env python3
"""生成《视觉验收标准说明》到桌面。"""
import json
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

REPORT_JSON = os.path.join(os.path.dirname(__file__), "visual_fidelity_report.json")
OUT = r"C:\Users\paz\Desktop\1212视觉验收标准.docx"


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    r.font.size = Pt(11)
    r.bold = bold
    para.paragraph_format.line_spacing = 1.35
    return para


def main():
    with open(REPORT_JSON, encoding="utf-8") as f:
        rep = json.load(f)

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("1212 视觉验收标准（取代字数判断）")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = "黑体"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    p(doc, "你说得对：此前用「字数 93%」「关键词全有」当标准，会严重高估质量。人眼看到的是版式、分页、字号、留白、表格线——这些原来没量。")
    p(doc, "")
    p(doc, "新标准（已落地为脚本 visual_fidelity_standard.py，可每次转换后自动跑）：", bold=True)
    for line in [
        "1. 渲染相似度：原 PDF / WPS docx / 我们 docx 各自打印成图，逐页算结构 SSIM（边缘）与像素 SSIM。",
        "2. 版式结构：段落数、对齐方式、表格/图片、文字量，与 WPS 同页对比。",
        "3. 文本重合：去空白后看本页与 WPS 是否同一内容（不能只看全文字数）。",
        "4. 硬性否决：整页大图、几乎无字、任一渲染相似度低于阈值。",
    ]:
        p(doc, line)

    p(doc, "")
    p(doc, f"当前样张实测：{rep['summary']['verdict']}，达标页 {rep['summary']['pass_ratio']:.0%}（{rep['summary']['total_pages']} 页中 {rep['summary']['pages_with_issues']} 页有问题）", bold=True)
    p(doc, "这与您肉眼感受一致：和原 PDF、WPS 范本差都很大；比「刚才那版」有进步，但远不能验收。")

    p(doc, "")
    p(doc, "阈值（可调）：")
    for k, v in rep["thresholds"].items():
        p(doc, f"  • {k} = {v}")

    p(doc, "")
    p(doc, "逐页结果摘要：", bold=True)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = ["页", "vs原PDF结构", "vs WPS结构", "版式", "文字重合", "主要问题"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    for pg in rep["pages"]:
        row = table.add_row().cells
        row[0].text = str(pg["page"])
        row[1].text = f"{pg.get('pdf_edge_ssim') or pg.get('pdf_ssim') or '-'}"
        row[2].text = f"{pg.get('wps_edge_ssim') or pg.get('wps_ssim') or '-'}"
        row[3].text = f"{pg.get('layout_vs_wps', '-')}"
        row[4].text = f"{pg.get('text_vs_wps', '-')}"
        row[5].text = "; ".join(pg.get("issues", [])[:2])

    p(doc, "")
    p(doc, "改进方向（由新标准反推，不是猜）：", bold=True)
    for line in [
        "• 像素/结构相似度低 → 需要「底图+文字」或更准的坐标排版，不能只堆可编辑字。",
        "• 与 WPS 版式分低 → 对齐、段距、字号、分页方式要向 WPS 靠拢。",
        "• 文字重合低 → 不是滤太多就是 OCR 块切分与 WPS 不一致。",
        "• 以后我说「可以验收」前，必须先跑本标准且 pass_ratio ≥ 75%。",
    ]:
        p(doc, line)

    p(doc, "")
    p(doc, "文件：1212验收标准报告.txt（明细）｜visual_fidelity_standard.py（脚本）")
    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    main()