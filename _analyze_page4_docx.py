#!/usr/bin/env python3
import zipfile
from docx import Document
from docx.enum.section import WD_ORIENT

path = r"C:\Users\paz\Desktop\P T W 测试\page_4_out.docx"
d = Document(path)
s = d.sections[0]
print("sections", len(d.sections), "portrait" if s.orientation == WD_ORIENT.PORTRAIT else "landscape")
print("size", round(s.page_width.inches, 2), "x", round(s.page_height.inches, 2))
print("paras", len([p for p in d.paragraphs if p.text.strip()]))
print("tables", len(d.tables))
body = "".join(p.text for p in d.paragraphs)
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            body += cell.text
print("chars", len(body))
print("head", body[:800])
with zipfile.ZipFile(path) as z:
    xml = z.read("word/document.xml").decode("utf-8")
print("anchors", xml.count("wp:anchor"))
if d.tables:
    t = d.tables[0]
    print("table", len(t.rows), "x", len(t.columns))