"""文件文字提取与翻译输出。"""
import os
import subprocess
import tempfile

from docx import Document


def extract_text_from_file(path, ext, ocr_pdf_func):
    ext = ext.lower()
    if ext == 'txt':
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().strip()
    if ext in ('doc', 'docx'):
        return _extract_docx(path)
    if ext == 'pdf':
        return _extract_pdf(path, ocr_pdf_func)
    raise ValueError('暂不支持该格式，请上传 TXT、Word 或 PDF')


def _open_docx(path):
    """打开Word文档，遇到老版.doc或损坏文件时给出明确提示，而不是把python-docx
    内部的报错原样抛给用户。"""
    with open(path, 'rb') as f:
        head = f.read(8)
    if head[:4] == b'\xd0\xcf\x11\xe0':
        raise ValueError('这是老版 .doc（97-2003）格式，暂不支持直接处理，请用Word打开后"另存为" .docx 格式再上传')
    try:
        return Document(path)
    except Exception:
        raise ValueError('无法读取这个Word文档，文件可能已损坏或不是标准.docx格式，请确认后重新上传')


def _extract_docx(path):
    doc = _open_docx(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not parts:
        raise ValueError('文档里没有读到文字')
    return '\n\n'.join(parts)


def _extract_pdf(path, ocr_pdf_func):
    result = subprocess.run(
        ['pdftotext', path, '-'],
        capture_output=True, text=True, timeout=60,
    )
    text = result.stdout.strip()
    if len(text) > 30:
        return text
    return ocr_pdf_func(path)


def _run_has_drawing(run):
    from docx.oxml.ns import qn
    el = run._element
    return bool(el.findall(qn('w:drawing'))) or bool(el.findall(qn('w:pict')))


def _replace_paragraph_text(paragraph, new_text):
    """把译文整段写进第一个文字run，清空其余文字run——保住这一段的主样式
    （字体/加粗/颜色/标题级别）。段落里若混了图片，图片所在的run跳过不动，
    不会被清空导致图片丢失；段落内局部混合格式（比如一句话中间一个词单独
    的颜色）会统一成第一个run的样式，这种情况比较少见，可以接受。"""
    text_runs = [r for r in paragraph.runs if not _run_has_drawing(r)]
    if not text_runs:
        return
    text_runs[0].text = new_text
    for r in text_runs[1:]:
        r.text = ''


def _iter_text_paragraphs(container):
    """递归遍历正文段落和表格单元格段落（含嵌套表格），跳过空段落和纯图片段落。"""
    for p in container.paragraphs:
        if p.text.strip():
            yield p
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_text_paragraphs(cell)


def translate_docx_inplace(input_path, output_path, translate_func, max_workers=6):
    """逐段翻译并原地写回原文档，保留段落样式、字体、表格、图片结构。
    translate_func(text) -> translated_text
    """
    from concurrent.futures import ThreadPoolExecutor

    doc = _open_docx(input_path)
    targets = list(_iter_text_paragraphs(doc))
    if not targets:
        raise ValueError('文档里没有读到文字')

    texts = [p.text for p in targets]
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        translated = list(ex.map(translate_func, texts))

    for p, new_text in zip(targets, translated):
        _replace_paragraph_text(p, new_text)

    doc.save(output_path)


def write_translated_docx(text, output_path):
    doc = Document()
    for block in text.split('\n\n'):
        if block.strip():
            doc.add_paragraph(block.strip())
    doc.save(output_path)


def write_ocr_docx(page_texts, output_path):
    """扫描件 OCR 结果写入 Word，保留换行，页间分页。"""
    doc = Document()
    for i, text in enumerate(page_texts):
        if i > 0:
            doc.add_page_break()
        if not text or not text.strip():
            continue
        for line in text.split('\n'):
            line = line.strip()
            if line:
                doc.add_paragraph(line)
    doc.save(output_path)


def translated_to_output(text, output_dir, uid, out_format='txt'):
    if out_format == 'docx':
        out = os.path.join(output_dir, f'{uid}.docx')
        write_translated_docx(text, out)
        return out, '译文.docx'
    out = os.path.join(output_dir, f'{uid}.txt')
    with open(out, 'w', encoding='utf-8') as f:
        f.write(text)
    return out, '译文.txt'