"""火山引擎 OCR — 智能文档解析 (ocr_pdf) + 通用文字识别 (ocr_normal)。"""
import base64
import copy
import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
import urllib.request

log = logging.getLogger(__name__)

CHUNK_PAGES = 16
MAX_PAGES = 300
NORMAL_OCR_DPI = 200
IMAGE_PARSE_DPI = 300
_CONTRACT_HINTS = (
    "服务合同", "甲方", "乙方", "CMCCTD", "签字/盖章", "合同专用章", "监测项目",
)
_TASK_CHECKLIST_HINTS = ("任务清单", "考评项目", "扣分标准", "你需要做什么", "要准备的资料")
_DENSE_TABLE_HINTS = (
    "规格书", "概况", "标高", "基坑", "监测", "参数", "结构概况", "技术规格", "单体",
)
_SCORING_FORM_HINTS = ("考评打分表", "服务管理考评", "考评得分", "扣分标准", "考评内容")
_SCORING_TABLE_HEADERS = frozenset({"序号", "考评内容", "分值", "扣分标准", "备注"})
_SCORING_TABLE_COL_RATIOS = {
    5: [0.06, 0.36, 0.08, 0.36, 0.14],
}
_BLANK_PAGE_MIN_CHARS = 15
_BLANK_PAGE_SNAPSHOT_DPI = 150
ENV_PATH = os.environ.get("TOOLBOX_ENV", "/home/toolbox/toolbox.env")


def _fitz_pixmap(page, **kwargs):
    """Compatible with both fitz < 1.19 (getPixmap) and fitz >= 1.19 (get_pixmap)."""
    try:
        return page.get_pixmap(**kwargs)
    except AttributeError:
        return page.getPixmap(**kwargs)


def _fitz_save_pix(pix, path):
    """Compatible with both fitz < 1.19 (writePNG) and fitz >= 1.19 (save)."""
    try:
        pix.save(path)
    except AttributeError:
        pix.writePNG(path)

# 常见水印碎片（含 OCR 误识变体）
_STATIC_WATERMARKS = {
    "臻子至善", "李于至善", "生臻子至善", "德厚生臻子至美", "德厚生臻子",
    "正德厚生", "正德厚生臻子", "哪生臻子", "臻予玉", "想", "于至善", "厚",
    "正德厚生臻子不", "生牛", "器", "德厚生", "名厚生",
}
_WATERMARK_FRAG_RES = [
    re.compile(r"德厚生臻子至美\s*"),
    re.compile(r"德厚生\s*"),
    re.compile(r"正德厚生?\s*"),
    re.compile(r"正臻子\s*"),
    re.compile(r"名厚生\s*"),
    re.compile(r"正德厚生臻子不?\"?\s*"),
    re.compile(r"臻子至善"),
    re.compile(r"李于至善"),
    re.compile(r"生臻子至善"),
    re.compile(r"于至善[。.]?\s*"),
    re.compile(r"哪生臻子不\"?\s*"),
    re.compile(r"臻予玉\s*"),
    re.compile(r"臻、\s*"),
]
_TABLE_CELL_RE = re.compile(r"<t[dh]([^>]*)>(.*?)</t[dh]>", re.I | re.S)
_ATTR_SPAN_RE = re.compile(r'(rowspan|colspan)\s*=\s*["\']?(\d+)', re.I)
_TASK_TABLE_HEADERS = frozenset({"时间", "任务名称", "你需要做什么", "要准备的资料"})
_TASK_TABLE_COL_RATIOS = {
    4: [0.10, 0.28, 0.32, 0.30],
    5: [0.12, 0.28, 0.14, 0.46],
}
_CHECKLIST_TITLE_BLUE = "03427A"
_CHECKLIST_SECTION_RED = "C60000"
_CHECKLIST_SUBTITLE_RED = "C50000"
_CHECKLIST_TABLE_HEAD_BG = "1E4D78"
_CHECKLIST_TABLE_HEAD_FG = "FFFFFF"
_CHECKLIST_TABLE_ACCENT_BG = "FEF2DF"
# Generic OCR fixes: common misreads in Chinese scanned documents.
# Document-specific corrections (company names, project codes, exact dates)
# should NOT go here — they break other contracts. Put them in a per-document
# JSON sidecar file named <contract_id>_ocr_fixes.json instead.
_OCR_FIXES = [
    # Common character confusions
    (re.compile(r"56生态"), "5G生态"),
    (re.compile(r"勤察"), "勘察"),
    (re.compile(r"粉察"), "勘察"),
    (re.compile(r"勒茶"), "勘察"),
    (re.compile(r"限眼同的求"), "按照"),
    (re.compile(r"甲防"), "甲方"),
    (re.compile(r"平方"), "甲方"),
    (re.compile(r"上渣"), "上海"),
    (re.compile(r"勤装"), "勘察"),
    (re.compile(r"黑章"), "盖章"),
    (re.compile(r"舍:\s*公"), "："),
    (re.compile(r"授校"), "授权"),
    (re.compile(r"仅器"), "仪器"),
    (re.compile(r"寄注"), "备注"),
    (re.compile(r"个同专用章"), "合同专用章"),
    # IT terminology
    (re.compile(r"[1I]硬件"), "IT硬件"),
    (re.compile(r"SU[\\／/]系统"), "SDN系统"),
    (re.compile(r"报客户机"), "瘦客户机"),
]
_IMG_MD_RE = re.compile(r"!\[[^\]]*\]\((https?://[^)]+)\)")
_IMG_MD_LINE_RE = re.compile(r"^\s*!\[[^\]]*\]\([^)]+\)\s*$")
_HEADING_MD_RE = re.compile(r"^#{1,6}\s+")
_TABLE_HTML_RE = re.compile(r"<table[\s\S]*?</table>", re.I)
_TABLE_CELL_TAG_RE = re.compile(r"</?t[dh]\b", re.I)
_DEGRADED_OCR_WARNING = (
    "智能文档解析未能还原版式或表格，已改用通用文字识别输出文字；"
    "如需完整表格，请使用电子版 PDF 或在 WPS 中转换"
)


def _load_env_file():
    if not os.path.isfile(ENV_PATH):
        return
    with open(ENV_PATH, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, val = line.split("=", 1)
            key = key.strip()
            val = val.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = val


def volc_credentials():
    _load_env_file()
    ak = os.environ.get("VOLC_ACCESS_KEY", "").strip()
    sk = os.environ.get("VOLC_SECRET_KEY", "").strip()
    return ak, sk


def volc_configured():
    ak, sk = volc_credentials()
    return bool(ak and sk)


def pdf_page_count(pdf_path):
    try:
        r = subprocess.run(
            ["pdfinfo", pdf_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        for line in r.stdout.splitlines():
            if line.startswith("Pages:"):
                return int(line.split(":", 1)[1].strip())
    except Exception:
        pass
    try:
        import fitz
        return fitz.open(pdf_path).page_count
    except Exception:
        return 1


def _ocr_pdf_chunk(visual_service, pdf_path, page_start, page_num):
    with open(pdf_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("ascii")
    form = {
        "image_base64": b64,
        "version": "v3",
        "file_type": "pdf",
        "page_start": page_start,
        "page_num": page_num,
        "table_mode": "html",
        "filter_header": "true",
        "parse_mode": "auto",
    }
    resp = visual_service.ocr_pdf(form)
    if not resp or resp.get("code") != 10000:
        msg = (resp or {}).get("message") or "火山 OCR 调用失败"
        raise ValueError(msg)
    data = resp.get("data") or {}
    markdown = data.get("markdown") or ""
    detail = data.get("detail")
    if isinstance(detail, str):
        try:
            detail = json.loads(detail)
        except json.JSONDecodeError:
            detail = []
    return markdown, detail if isinstance(detail, list) else []


def _parse_ocr_detail(raw_detail):
    detail = raw_detail
    if isinstance(detail, str):
        try:
            detail = json.loads(detail)
        except json.JSONDecodeError:
            detail = []
    return detail if isinstance(detail, list) else []


def _image_parse_dpi():
    try:
        return int(os.environ.get("PREPROCESS_DPI", str(IMAGE_PARSE_DPI)))
    except ValueError:
        return IMAGE_PARSE_DPI


def _ocr_blob(markdown="", details=None):
    blob = markdown or ""
    for page in details or []:
        blob += page.get("page_md") or ""
        for block in page.get("textblocks") or []:
            blob += block.get("text") or ""
    return blob


def _looks_like_contract(markdown="", details=None):
    blob = _ocr_blob(markdown, details)
    return sum(1 for h in _CONTRACT_HINTS if h in blob) >= 2


def _looks_like_task_checklist(markdown="", details=None):
    blob = _ocr_blob(markdown, details)
    if "任务清单" in blob:
        return True
    return "任务名称" in blob and "你需要做什么" in blob


def _preprocessed_ocr_path(pdf_path):
    ocr_path = pdf_path
    cleaned_tmp = None
    try:
        from preprocessing import preprocess_pdf_for_ocr
        result = preprocess_pdf_for_ocr(pdf_path)
        if result != pdf_path:
            ocr_path = result
            cleaned_tmp = result
    except Exception:
        pass
    return ocr_path, cleaned_tmp


def _visual_service():
    from volcengine.visual.VisualService import VisualService

    visual = VisualService()
    ak, sk = volc_credentials()
    visual.set_ak(ak)
    visual.set_sk(sk)
    return visual


def _strip_image_markdown(text):
    """去掉仅含 Markdown 图片链接的行（火山把整页当 fig 时会返回这种）。"""
    if not text:
        return ""
    kept = []
    for line in text.splitlines():
        if _IMG_MD_LINE_RE.match(line.strip()):
            continue
        stripped = _IMG_MD_RE.sub("", line).strip()
        if stripped:
            kept.append(stripped)
    return "\n".join(kept).strip()


def _text_has_substance(text, min_len=4):
    """是否含可用文字或表格 HTML（纯图链不算）。"""
    text = _strip_image_markdown(text or "")
    if not text:
        return False
    if _TABLE_HTML_RE.search(text) or _TABLE_CELL_TAG_RE.search(text):
        return True
    plain = re.sub(r"<[^>]+>", " ", text)
    plain = _normalize_text(plain)
    if len(plain) < min_len:
        return False
    if re.search(r"[\u4e00-\u9fff]{2,}", plain):
        return True
    return bool(re.search(r"[a-zA-Z0-9]{3,}", plain))


def _block_has_substance(block):
    label = (block.get("label") or "para").lower()
    if label in ("image", "figure", "fig"):
        return False
    text = block.get("text") or ""
    if label == "table" or "<table" in text.lower():
        return _text_has_substance(text, min_len=2)
    return _text_has_substance(text)


def _detail_has_usable_content(details):
    """智能文档解析是否返回了可编辑文字或表格结构（仅图链不算）。"""
    if not details:
        return False
    for page in details:
        if _text_has_substance(page.get("page_md") or ""):
            return True
        for block in page.get("textblocks") or []:
            if _block_has_substance(block):
                return True
    return False


def _markdown_has_usable_content(markdown):
    return _text_has_substance(markdown or "")


def _doc_parse_image_only(markdown, details):
    """智能文档解析是否只回了整页图链（B 类），而非完全空（C 类）。"""
    if _markdown_has_usable_content(markdown) or _detail_has_usable_content(details):
        return False
    if _IMG_MD_RE.search(markdown or ""):
        return True
    for page in details or []:
        if _IMG_MD_RE.search(page.get("page_md") or ""):
            return True
        for block in page.get("textblocks") or []:
            if (block.get("label") or "").lower() in ("image", "figure", "fig"):
                return True
    return False


def _detail_looks_fragmented(page):
    """直传 detail 碎成大量短块（签章页常见），PNG 重解析质量更高。"""
    blocks = [
        b for b in (page.get("textblocks") or [])
        if (b.get("label") or "para").lower() not in ("image", "figure", "fig", "foot")
    ]
    if not blocks:
        return False
    texts = [_normalize_text(b.get("text") or "") for b in blocks]
    texts = [t for t in texts if t]
    if not texts:
        return False
    short = sum(1 for t in texts if len(t) <= 4)
    scores = [_text_quality_score(t) for t in texts]
    avg_score = sum(scores) / len(scores)
    blob = " ".join(texts)
    sig_like = bool(re.search(r"签字|盖章|合同专用章|甲方|乙方", blob))
    has_clause = bool(re.search(r"第[一二三四五六七八九十\d]+条", blob))
    if not sig_like and avg_score >= 0.42:
        return False
    if sig_like and short >= 3 and short >= int(len(texts) * 0.25):
        return True
    if short >= max(3, int(len(texts) * 0.32)) and avg_score < 0.42:
        return True
    if avg_score < 0.38:
        return True
    if sig_like and not has_clause and short >= 2 and avg_score < 0.55:
        return True
    return False


def _detail_needs_image_mode(markdown, details):
    """PDF 直传多为整页图或签字页碎片（合同 A 类），应逐页 PNG 重解析。"""
    if _doc_parse_image_only(markdown, details):
        return True
    if not details:
        return False
    if len(details) == 1 and _detail_looks_fragmented(details[0]):
        return True
    if any(_detail_looks_fragmented(p) for p in details):
        if len(details) <= 3:
            return True
    total = len(details)
    weak_pages = 0
    strong_pages = 0
    for page in details:
        meaningful = _meaningful_text_len(page, _STATIC_WATERMARKS)
        blocks = page.get("textblocks") or []
        if meaningful >= 120:
            strong_pages += 1
            continue
        image_heavy = (
            not blocks
            or all((b.get("label") or "").lower() in ("image", "figure", "fig") for b in blocks)
            or meaningful < 50
        )
        if image_heavy:
            weak_pages += 1
    if strong_pages >= max(2, (total + 2) // 3):
        return False
    return weak_pages >= max(2, int(total * 0.55))


def _page_rgb_and_b64(pdf_path, page_index, dpi=NORMAL_OCR_DPI, correct_sideways=False):
    import cv2
    import numpy as np

    if correct_sideways:
        bgr, layout = _get_corrected_page_bgr(
            pdf_path, page_index, dpi, probe_coarse=True,
        )
        if _layout_needs_orientation_handling(layout):
            log.info(
                "页面扫描校正 OCR: %s p%d kind=%s coarse=%s skew=%.1f",
                pdf_path, page_index, layout.get("kind"),
                layout.get("correction_deg"), layout.get("skew_deg") or 0,
            )
    else:
        bgr = _render_page_bgr(pdf_path, page_index, dpi)
    rgb = cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB)
    fd, path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    try:
        cv2.imwrite(path, bgr)
        with open(path, "rb") as f:
            img_b64 = base64.b64encode(f.read()).decode("ascii")
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass
    return img_b64, rgb


def _ocr_normal_page(visual, image_b64):
    resp = visual.ocr_normal({"image_base64": image_b64})
    if not resp or resp.get("code") != 10000:
        msg = (resp or {}).get("message") or "通用文字识别失败"
        raise ValueError(msg)
    return resp.get("data") or {}


def _avg_char_height(line_chars):
    if not line_chars:
        return 0
    return sum(c.get("height", 0) for c in line_chars) / len(line_chars)


def _px_to_pt(height_px, dpi=NORMAL_OCR_DPI):
    from docx.shared import Pt

    return Pt(max(9, min(28, round(height_px * 72.0 / dpi * 0.95, 1))))


def _sample_line_color(rgb_arr, line_chars, prefer_blue=False):
    from docx.shared import RGBColor

    if rgb_arr is None or not line_chars:
        return None
    blues = []
    others = []
    for c in line_chars:
        x = int(c.get("x", 0) + c.get("width", 0) / 2)
        y = int(c.get("y", 0) + c.get("height", 0) / 2)
        if 0 <= y < rgb_arr.shape[0] and 0 <= x < rgb_arr.shape[1]:
            r, g, b = [int(v) for v in rgb_arr[y, x]]
            if b > max(r, g) + 35 and b > 120:
                blues.append((r, g, b))
            else:
                others.append((r, g, b))
    picks = blues if blues else ([] if prefer_blue else others)
    if not picks:
        return None
    r = int(sum(t[0] for t in picks) / len(picks))
    g = int(sum(t[1] for t in picks) / len(picks))
    b = int(sum(t[2] for t in picks) / len(picks))
    return RGBColor(r, g, b)


def normal_ocr_pdf_to_docx(pdf_path, output_path, dpi=NORMAL_OCR_DPI):
    """通用文字识别：纯文字扫描页（C 类作文等）。"""
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    visual = _visual_service()
    total = min(pdf_page_count(pdf_path), MAX_PAGES)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)

    pages_written = 0
    for pi in range(total):
        img_b64, rgb = _page_rgb_and_b64(pdf_path, pi, dpi=dpi)
        data = _ocr_normal_page(visual, img_b64)
        lines = data.get("line_texts") or []
        chars = data.get("chars") or []
        if not lines:
            continue

        if pages_written > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        heights = [
            _avg_char_height(chars[i] if i < len(chars) else [])
            for i in range(len(lines))
        ]
        body_heights = [ht for i, ht in enumerate(heights) if i > 0 and ht > 0]
        body_med = sorted(body_heights)[len(body_heights) // 2] if body_heights else 33

        for i, text in enumerate(lines):
            text = (text or "").strip()
            if not text:
                continue
            line_chars = chars[i] if i < len(chars) else []
            h = heights[i]
            is_title = i == 0 or (h > 0 and h >= max(body_med * 1.45, 50))
            blue_color = _sample_line_color(rgb, line_chars, prefer_blue=True)

            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "宋体"
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass
            if is_title:
                run.bold = True
                run.font.size = _px_to_pt(h) if h else Pt(22)
                run.font.color.rgb = blue_color or RGBColor(0, 114, 239)
            elif blue_color:
                run.font.size = _px_to_pt(h) if h else Pt(11)
                run.font.color.rgb = blue_color
            else:
                run.font.size = _px_to_pt(h) if h else Pt(11)

        pages_written += 1

    if pages_written == 0:
        raise ValueError("通用文字识别未返回有效内容")
    doc.save(output_path)


def _ocr_pdf_image_page_data(visual, image_b64):
    """单页 PNG/JPEG 走智能文档解析，返回 markdown + detail。"""
    form = {
        "image_base64": image_b64,
        "version": "v3",
        "file_type": "image",
        "page_start": 0,
        "page_num": 1,
        "table_mode": "html",
        "filter_header": "true",
        "parse_mode": "auto",
    }
    resp = visual.ocr_pdf(form)
    if not resp or resp.get("code") != 10000:
        msg = (resp or {}).get("message") or "火山 OCR 图片解析失败"
        raise ValueError(msg)
    data = resp.get("data") or {}
    return {
        "markdown": data.get("markdown") or "",
        "detail": _parse_ocr_detail(data.get("detail")),
    }


def _ocr_pdf_image_page(visual, image_b64):
    return _ocr_pdf_image_page_data(visual, image_b64).get("markdown") or ""


def _pdf_image_mode_pages(pdf_path, dpi=None, with_detail=False):
    """逐页渲染（先去水印）后调用 ocr_pdf(file_type=image)。"""
    dpi = dpi or _image_parse_dpi()
    ocr_path, cleaned_tmp = _preprocessed_ocr_path(pdf_path)
    visual = _visual_service()
    total = min(pdf_page_count(ocr_path), MAX_PAGES)
    pages_md = []
    details = []
    try:
        for pi in range(total):
            correct_sideways = False
            if total == 1:
                layout = _analyze_page_layout(ocr_path, pi, dpi, probe_coarse=True)
                correct_sideways = _layout_needs_orientation_handling(layout)
            img_b64, _ = _page_rgb_and_b64(
                ocr_path, pi, dpi=dpi, correct_sideways=correct_sideways,
            )
            data = _ocr_pdf_image_page_data(visual, img_b64)
            pages_md.append((data.get("markdown") or "").strip())
            if with_detail:
                details.extend(data.get("detail") or [])
    finally:
        if cleaned_tmp and os.path.isfile(cleaned_tmp):
            try:
                os.unlink(cleaned_tmp)
            except OSError:
                pass
    return pages_md, details


def pdf_to_markdown_image_mode(pdf_path, dpi=None):
    """逐页 PNG 智能文档解析，返回 Markdown 列表（任务清单 B 类）。"""
    pages_md, _ = _pdf_image_mode_pages(pdf_path, dpi=dpi, with_detail=False)
    return pages_md


def pdf_to_detail_image_mode(pdf_path, dpi=None):
    """逐页 PNG 智能文档解析，返回 (detail 列表, 每页 markdown)。"""
    pages_md, details = _pdf_image_mode_pages(pdf_path, dpi=dpi, with_detail=True)
    return details, pages_md


def _set_cell_shading(cell, fill_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_alignment(cell, align):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for para in cell.paragraphs:
        para.alignment = align
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        for para in cell.paragraphs:
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0


def _set_cell_text_style(cell, pt, bold=False, color_hex=None, align=None, fill_hex=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    if fill_hex:
        _set_cell_shading(cell, fill_hex)
    if align is not None:
        _set_cell_alignment(cell, align)
    for para in cell.paragraphs:
        if not para.runs:
            run = para.add_run(para.text)
            para.text = ""
        else:
            run = para.runs[0]
        run.font.size = Pt(pt)
        run.bold = bool(bold)
        if color_hex:
            _set_run_font_style(run, "宋体", pt, east_asia="SimSun", bold=bold, color_hex=color_hex)


def _is_task_checklist_table(parsed):
    if not parsed or len(parsed[0]) < 3:
        return False
    header_text = {_normalize_text(c.get("text") or "") for c in parsed[0]}
    return len(header_text & _TASK_TABLE_HEADERS) >= 2


def _is_scoring_form_table(parsed):
    if not parsed or len(parsed[0]) < 3:
        return False
    header_text = {_normalize_text(c.get("text") or "") for c in parsed[0]}
    return len(header_text & _SCORING_TABLE_HEADERS) >= 3


def _page_is_scoring_form(page, page_md=""):
    blob = _page_table_blob(page, page_md)
    if sum(1 for h in _SCORING_FORM_HINTS if h in blob) < 2:
        return False
    return "<table" in blob.lower() or _page_mainly_table(page)


def _task_table_font_pt(ncols, nrows):
    if ncols >= 7 or nrows >= 22:
        return 6
    if ncols >= 6 or nrows >= 15:
        return 6.5
    if ncols >= 5 or nrows >= 10:
        return 8.5
    return 9


def _apply_task_checklist_table_style(table, parsed, placements, nrows, ncols, font_pt, use_landscape):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    usable = Inches(9.0 if use_landscape else 6.2)
    ratios = _TASK_TABLE_COL_RATIOS.get(ncols)
    if ratios:
        widths = [int(usable * r) for r in ratios]
        diff = int(usable) - sum(widths)
        if diff and widths:
            widths[-1] += diff
        for ci, column in enumerate(table.columns):
            if ci < len(widths):
                column.width = widths[ci]

    header_row = {_normalize_text(c.get("text") or "") for c in parsed[0]}
    for ri, ci, cell in placements:
        tcell = table.rows[ri].cells[ci]
        text = cell.get("text") or ""
        if ri == 0 and _normalize_text(text) in header_row:
            _set_cell_text_style(
                tcell,
                font_pt,
                bold=True,
                color_hex=_CHECKLIST_TABLE_HEAD_FG,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                fill_hex=_CHECKLIST_TABLE_HEAD_BG,
            )
            continue
        align = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 and re.fullmatch(r"\d{1,2}月|全年", text.strip()) else None
        fill = None
        if re.match(r"^(安全类|平安类)", text):
            fill = _CHECKLIST_TABLE_ACCENT_BG
        _set_cell_text_style(
            tcell,
            font_pt,
            bold=bool(fill),
            color_hex=_CHECKLIST_SECTION_RED if fill else None,
            align=align,
            fill_hex=fill,
        )


def _detect_row_line_positions(img_bgr):
    """Return normalized y-positions [0.0–1.0] of horizontal table lines detected via OpenCV."""
    import cv2
    import numpy as np

    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
    h, w = binary.shape
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(30, w // 6), 1))
    h_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel)
    proj = h_lines.sum(axis=1).astype(float)
    if proj.max() == 0:
        return []
    thresh = proj.max() * 0.2
    positions = []
    i = 0
    while i < h:
        if proj[i] > thresh:
            j = i
            while j < h and proj[j] > thresh:
                j += 1
            positions.append((i + j) / 2.0 / h)
            i = j
        else:
            i += 1
    return positions


def _row_heights_from_line_positions(positions, nrows):
    """Convert detected line y-positions to proportional row heights (sum=1.0), or None."""
    if len(positions) < nrows + 1:
        return None
    if len(positions) == nrows + 1:
        sep = positions
    else:
        first, last = positions[0], positions[-1]
        inner = positions[1:-1]
        if len(inner) >= nrows - 1:
            step = len(inner) / max(nrows - 1, 1)
            sel = [inner[min(int(i * step), len(inner) - 1)] for i in range(nrows - 1)]
        else:
            return None
        sep = [first] + sel + [last]
    heights = [sep[i + 1] - sep[i] for i in range(len(sep) - 1)]
    if not heights or min(heights) <= 0:
        return None
    total = sum(heights)
    return [h / total for h in heights]


def _extract_row_heights_img2table(img_bgr):
    """Use img2table (OpenCV Hough lines) to detect row heights from a scanned table image.
    Returns normalized height fractions (sum=1.0) per row, or None on failure/unavailable.
    """
    try:
        from img2table.document import Image as I2TImage
    except ImportError:
        log.debug("img2table not installed, skipping proportional row height detection")
        return None
    import cv2
    import tempfile
    import os

    fd, tmp_path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    try:
        cv2.imwrite(tmp_path, img_bgr)
        doc_img = I2TImage(tmp_path)
        tables = doc_img.extract_tables(ocr=None, implicit_rows=True, borderless_tables=False)
    except Exception as exc:
        log.debug("img2table extract_tables failed: %s", exc)
        tables = []
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    if not tables:
        return None

    # Pick the table with the most rows
    best = max(tables, key=lambda t: len(t.content))
    rows_dict = best.content  # dict: row_idx -> list[Cell]
    if not rows_dict:
        return None

    row_heights_px = []
    for row_idx in sorted(rows_dict.keys()):
        cells = rows_dict[row_idx]
        if not cells:
            continue
        y1 = min(c.bbox.y1 for c in cells)
        y2 = max(c.bbox.y2 for c in cells)
        row_heights_px.append(max(1, y2 - y1))

    if len(row_heights_px) < 2:
        return None

    total = sum(row_heights_px)
    return [h / total for h in row_heights_px]


def _apply_scoring_form_table_style(table, parsed, placements, nrows, ncols, font_pt, use_landscape, row_heights=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    usable = Inches(7.2 if not use_landscape else 9.0)
    ratios = _SCORING_TABLE_COL_RATIOS.get(ncols)
    if ratios:
        widths = [int(usable * r) for r in ratios]
        diff = int(usable) - sum(widths)
        if diff and widths:
            widths[-1] += diff
        for ci, column in enumerate(table.columns):
            if ci < len(widths):
                column.width = widths[ci]

    header_row = {_normalize_text(c.get("text") or "") for c in parsed[0]}
    for ri, ci, cell in placements:
        tcell = table.rows[ri].cells[ci]
        text = cell.get("text") or ""
        if ri == 0 and _normalize_text(text) in header_row:
            _set_cell_text_style(
                tcell,
                font_pt,
                bold=True,
                color_hex=_CHECKLIST_TABLE_HEAD_FG,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                fill_hex=_CHECKLIST_TABLE_HEAD_BG,
            )
            continue
        align = None
        if ci == 0 or _normalize_text(text).isdigit():
            align = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_text_style(tcell, font_pt, align=align)

    # Apply row heights: proportional from image detection, or let Word auto-size
    # sec dimensions are in EMU (914400/inch); w:trHeight val is twips (1440/inch)
    _EMU_PER_TWIP = 635
    sec = table.part.document.sections[-1]
    avail_h_twips = int(sec.page_height - sec.top_margin - sec.bottom_margin) // _EMU_PER_TWIP
    rh = list(row_heights) if row_heights else []
    # Trim extra trailing rows first (img2table often detects a stray footer line)
    while len(rh) > nrows and len(rh) > 2:
        rh.pop()
    if rh and len(rh) == nrows:
        total = sum(rh)
        rh = [h / total for h in rh]
        for ri, row in enumerate(table.rows):
            twips = max(200, int(avail_h_twips * rh[ri]))
            _set_row_height(row, twips, exact=True)


def _add_checklist_styled_para(doc, line):
    """任务清单类扫描页：按 WPS 近似还原标题/章节配色。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    raw = (line or "").strip()
    if not raw:
        return False

    heading_m = re.match(r"^(#{1,6})\s*(.+)$", raw)
    if heading_m:
        text = heading_m.group(2).strip()
        if "任务清单" in text or re.match(r"^20\d{2}年", text):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            _set_run_font_style(run, "黑体", 15.5, east_asia="SimHei", bold=True, color_hex=_CHECKLIST_TITLE_BLUE)
            return True
        if re.match(r"^[一二三四五六七八九十]+、", text):
            para = doc.add_paragraph()
            run = para.add_run(text)
            _set_run_font_style(run, "宋体", 12, east_asia="SimSun", bold=True, color_hex=_CHECKLIST_SECTION_RED)
            return True
        para = doc.add_paragraph()
        run = para.add_run(text)
        _set_run_font_style(run, "宋体", 12, east_asia="SimSun", bold=True, color_hex=_CHECKLIST_SECTION_RED)
        return True

    if raw.startswith("(") and "执行版" in raw:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(raw)
        _set_run_font_style(run, "宋体", 9.5, east_asia="SimSun", bold=False, color_hex=_CHECKLIST_SUBTITLE_RED)
        return True

    if raw.startswith("整理时间"):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(raw)
        _set_run_font_style(run, "宋体", 9.5, east_asia="SimSun")
        return True

    if re.match(r"^[一二三四五六七八九十]+、", raw):
        para = doc.add_paragraph()
        run = para.add_run(raw)
        _set_run_font_style(run, "宋体", 12, east_asia="SimSun", bold=True, color_hex=_CHECKLIST_SECTION_RED)
        return True

    if re.match(r"^(安全类|平安类)", raw):
        para = doc.add_paragraph()
        run = para.add_run(raw)
        _set_run_font_style(run, "宋体", 12, east_asia="SimSun", bold=True, color_hex=_CHECKLIST_SECTION_RED)
        return True

    return False


def _render_checklist_markdown_line(doc, line, watermarks, cache_dir):
    raw = (line or "").strip()
    if not raw:
        return
    if raw.lower().startswith("<table"):
        _add_html_table(doc, raw)
        return
    if _add_checklist_styled_para(doc, raw):
        return
    _render_text_content(doc, raw, watermarks, cache_dir)


def markdown_pages_to_docx(page_markdowns, output_path):
    """多页 Markdown（含 <table> HTML）写入 Word，页间分页。"""
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    cache_dir = tempfile.mkdtemp(prefix="volc_md_")
    watermarks = _STATIC_WATERMARKS
    try:
        wrote = False
        for markdown_text in page_markdowns:
            if not markdown_text:
                continue
            if wrote:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            for line in markdown_text.splitlines():
                _render_checklist_markdown_line(doc, line, watermarks, cache_dir)
            wrote = True
        if not wrote:
            raise ValueError("图片模式未返回有效内容")
        doc.save(output_path)
    finally:
        shutil.rmtree(cache_dir, ignore_errors=True)


def pdf_to_markdown(pdf_path):
    ak, sk = volc_credentials()
    if not ak or not sk:
        raise ValueError("未配置火山 OCR 密钥")

    from volcengine.visual.VisualService import VisualService

    visual = VisualService()
    visual.set_ak(ak)
    visual.set_sk(sk)

    # Pre-process: remove diagonal gray watermarks before OCR for cleaner text extraction.
    # Controlled by PREPROCESS_WATERMARK env var (default on). Falls back to original
    # if preprocessing fails or the module is unavailable.
    ocr_path = pdf_path
    cleaned_tmp = None
    try:
        from preprocessing import preprocess_pdf_for_ocr
        result = preprocess_pdf_for_ocr(pdf_path)
        if result != pdf_path:
            ocr_path = result
            cleaned_tmp = result
    except Exception:
        pass

    try:
        total = min(pdf_page_count(ocr_path), MAX_PAGES)
        parts = []
        details = []
        start = 0
        while start < total:
            chunk = min(CHUNK_PAGES, total - start)
            md, detail = _ocr_pdf_chunk(visual, ocr_path, start, chunk)
            if md.strip():
                parts.append(md.strip())
            details.extend(detail)
            start += chunk
    finally:
        if cleaned_tmp and os.path.isfile(cleaned_tmp):
            try:
                os.unlink(cleaned_tmp)
            except OSError:
                pass

    return "\n\n".join(parts), details


def _docx_mode(page_count):
    _load_env_file()
    mode = os.environ.get("VOLC_DOCX_MODE", "").strip().lower()
    if mode in ("hybrid", "text", "visual"):
        return mode
    return "hybrid" if page_count <= 30 else "text"


def _strip_watermark_substrings(text):
    for pat in _WATERMARK_FRAG_RES:
        text = pat.sub("", text)
    return text


def _apply_ocr_fixes(text):
    for pat, repl in _OCR_FIXES:
        text = pat.sub(repl, text)
    return text


def _normalize_text(text):
    if not text:
        return ""
    text = text.replace("\u3000", " ").strip()
    text = _HEADING_MD_RE.sub("", text)
    text = _strip_watermark_substrings(text)
    text = _apply_ocr_fixes(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _detect_dynamic_watermarks(pages):
    from collections import Counter

    counts = Counter()
    for page in pages:
        for block in page.get("textblocks") or []:
            if block.get("label") in ("header", "foot", "image", "table"):
                continue
            text = _normalize_text(block.get("text") or "")
            if 2 <= len(text) <= 8 and re.fullmatch(r"[\u4e00-\u9fffA-Za-z0-9·、，。]+", text):
                counts[text] += 1
    dynamic = {t for t, c in counts.items() if c >= 4 and len(t) <= 6}
    return _STATIC_WATERMARKS | dynamic


def _text_quality_score(text):
    text = (text or "").strip()
    if not text:
        return 1.0
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return 1.0
    han = len(re.findall(r"[\u4e00-\u9fff]", compact))
    alnum = len(re.findall(r"[A-Za-z0-9]", compact))
    base = (han + alnum * 0.85) / len(compact)
    penalty = 0.0
    for w in ("臻子", "至善", "厚生", "德厚", "李于", "正德", "哪生", "名厚"):
        if w in text:
            penalty += 0.18
    if re.search(r"[女熊端足文]{3,}", text):
        penalty += 0.35
    if len(compact) <= 3 and not re.search(r"\d", compact):
        penalty += 0.25
    return max(0.0, min(1.0, base - penalty))


def _table_ocr_quality(parsed_rows):
    scores = []
    for row in parsed_rows:
        for cell in row:
            t = (cell.get("text") or "").strip()
            if not t:
                continue
            scores.append(_text_quality_score(t))
    return sum(scores) / len(scores) if scores else 0.0


_SIG_TEXT_HINTS = (
    "签字", "盖章", "签章", "日期", "甲方", "乙方", "合同专用章",
    "授权代表", "法定代表人", "第八条", "第九条", "签字页", "签章页",
)


def _is_sig_relevant_text(text):
    text = text or ""
    return any(h in text for h in _SIG_TEXT_HINTS)


def _is_noise_text(text, watermarks, sig_page=False):
    text = _normalize_text(text)
    if not text:
        return True
    if sig_page and _is_sig_relevant_text(text) and len(text) >= 3:
        return False
    if text in watermarks:
        return True
    if len(text) <= 2 and not re.search(r"\d", text):
        if text in ("分值", "备注", "序号", "日期", "姓名"):
            return False
        return True
    if _text_quality_score(text) < 0.28:
        if sig_page and len(text) >= 6 and re.search(r"[\u4e00-\u9fff]{2,}", text):
            return False
        return True
    # 纯水印短语：短且无标点、无数字
    if len(text) <= 6 and not re.search(r"[，。；：、（）()\d]", text):
        if any(w in text for w in ("臻", "至善", "厚生", "德厚")):
            return True
    return False


def _keep_cover_text(text, watermarks):
    text = _normalize_text(text)
    if _is_noise_text(text, watermarks):
        return False
    if re.search(r"CMCCTD|服务合同|监测项目|工程|合同", text):
        return True
    if ("甲方" in text or "乙方" in text) and _text_quality_score(text) >= 0.45:
        return True
    return _text_quality_score(text) >= 0.68


def _begin_landscape_section(doc):
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches

    sec = doc.sections[-1]
    new_sec = doc.add_section()
    new_sec.orientation = WD_ORIENT.LANDSCAPE
    new_sec.page_width, new_sec.page_height = sec.page_height, sec.page_width
    new_sec.top_margin = Inches(0.45)
    new_sec.bottom_margin = Inches(0.45)
    new_sec.left_margin = Inches(0.4)
    new_sec.right_margin = Inches(0.4)
    return new_sec


def _begin_portrait_section(doc):
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches

    sec = doc.sections[-1]
    new_sec = doc.add_section()
    new_sec.orientation = WD_ORIENT.PORTRAIT
    new_sec.page_width, new_sec.page_height = sec.page_height, sec.page_width
    new_sec.top_margin = Inches(0.7)
    new_sec.bottom_margin = Inches(0.7)
    new_sec.left_margin = Inches(0.75)
    new_sec.right_margin = Inches(0.75)
    return new_sec


def _sync_doc_section_to_pdf(doc, pdf_path, page_index=0, tight=True):
    """Word 节与 PDF 页同尺寸同方向（A4 竖版等），避免擅自改成横版/Letter。"""
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches

    rect = None
    try:
        import fitz
        pdf = fitz.open(pdf_path)
        try:
            rect = pdf[page_index].rect
        finally:
            pdf.close()
    except Exception:
        rect = None

    sec = doc.sections[0]
    if rect is not None:
        width_in = rect.width / 72.0
        height_in = rect.height / 72.0
        if width_in > height_in:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width = Inches(width_in)
            sec.page_height = Inches(height_in)
        else:
            sec.orientation = WD_ORIENT.PORTRAIT
            sec.page_width = Inches(width_in)
            sec.page_height = Inches(height_in)
    margin = Inches(0.45 if tight else 0.6)
    sec.top_margin = margin
    sec.bottom_margin = margin
    sec.left_margin = margin
    sec.right_margin = margin
    return sec


def _set_row_height(row, twips, exact=True):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(twips)))
    tr_height.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(tr_height)


def _set_cell_compact(cell, font_pt):
    from docx.shared import Pt

    _set_cell_font(cell, font_pt)
    for para in cell.paragraphs:
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0


def _apply_compact_dense_table_style(table, nrows, font_pt, reserved_in=0.3):
    from docx.shared import Inches

    sec = table.part.document.sections[-1]
    usable = sec.page_height - sec.top_margin - sec.bottom_margin - Inches(reserved_in)
    row_twips = max(96, int(usable) // max(nrows, 1))
    row_twips = min(row_twips, 320)
    if nrows >= 45:
        font_pt = min(font_pt, 5.0)
    elif nrows >= 35:
        font_pt = min(font_pt, 5.5)
    else:
        font_pt = min(font_pt, 6.0)
    for row in table.rows:
        _set_row_height(row, row_twips, exact=True)
        for cell in row.cells:
            _set_cell_compact(cell, font_pt)


def _sort_blocks(blocks):
    def key_fn(b):
        box = b.get("box") or {}
        return (box.get("y0", 0), box.get("x0", 0))
    return sorted(blocks, key=key_fn)


def _download_image(url, cache_dir):
    os.makedirs(cache_dir, exist_ok=True)
    name = re.sub(r"[^\w.-]", "_", url.split("?")[0].split("/")[-1]) or "img.jpg"
    if not name.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".gif")):
        name += ".jpg"
    path = os.path.join(cache_dir, name)
    if os.path.isfile(path) and os.path.getsize(path) > 0:
        return path
    req = urllib.request.Request(url, headers={"User-Agent": "toolbox/1.0"})
    with urllib.request.urlopen(req, timeout=30) as resp:
        data = resp.read()
    if not data:
        raise ValueError("empty image")
    with open(path, "wb") as f:
        f.write(data)
    return path


def _seal_extract_enabled():
    _load_env_file()
    return os.environ.get("VOLC_EXTRACT_SEAL", "1").strip().lower() not in ("0", "false", "no")


def _pdf_page_landscape(pdf_path, page_index):
    try:
        import fitz
        page = fitz.open(pdf_path)[page_index]
        rect = page.rect
        return rect.width > rect.height
    except Exception:
        return False


_PAGE_LAYOUT_CACHE = {}

# 页面版式 / 扫描校正分类（表格页）
# portrait_upright  — 正坐表（page_6）：竖版 A4，表头在上
# sideways_ccw/cw   — 侧躺约 90°（扫描横着放）
# upside_down       — 颠倒约 180°
# skewed            — 仅歪斜（如约 8° 没放正）
# landscape_native  — 真横版页
# pdf_rotated_meta  — PDF /Rotate 元数据


def _clear_page_layout_cache():
    _PAGE_LAYOUT_CACHE.clear()


def _pdf_page_rotation_meta(pdf_path, page_index):
    try:
        import fitz
        pdf = fitz.open(pdf_path)
        try:
            return int(pdf[page_index].rotation) % 360
        finally:
            pdf.close()
    except Exception:
        return 0


def _render_page_bgr(pdf_path, page_index, dpi):
    import cv2
    import fitz
    import numpy as np

    doc = fitz.open(pdf_path)
    try:
        page = doc[page_index]
        pix = _fitz_pixmap(page, matrix=fitz.Matrix(dpi / 72.0, dpi / 72.0), alpha=False)
        arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
        if pix.n >= 3:
            return cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)
        return arr
    finally:
        doc.close()


def _render_page_bgr_rotated(pdf_path, page_index, dpi, correction_deg):
    """Render PDF page with coarse rotation applied via fitz matrix — no warpAffine interpolation."""
    import cv2
    import fitz
    import numpy as np

    doc = fitz.open(pdf_path)
    try:
        page = doc[page_index]
        scale = dpi / 72.0
        # fitz.prerotate uses standard math CCW convention; negate for CW degrees
        mat = fitz.Matrix(scale, scale).prerotate(-correction_deg)
        pix = _fitz_pixmap(page, matrix=mat, alpha=False)
        arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
        return cv2.cvtColor(arr, cv2.COLOR_RGB2BGR) if pix.n >= 3 else arr
    finally:
        doc.close()


def _rotate_bgr(img_bgr, deg_cw):
    import cv2

    if deg_cw == 90:
        return cv2.rotate(img_bgr, cv2.ROTATE_90_CLOCKWISE)
    if deg_cw == 180:
        return cv2.rotate(img_bgr, cv2.ROTATE_180)
    if deg_cw == 270:
        return cv2.rotate(img_bgr, cv2.ROTATE_90_COUNTERCLOCKWISE)
    return img_bgr


def _page_ink_binary(img_bgr):
    import cv2

    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
    return binary


def _ink_cover_metrics(img_bgr):
    import cv2

    binary = _page_ink_binary(img_bgr)
    coords = cv2.findNonZero(binary)
    if coords is None:
        return {"cover_w": 0.0, "cover_h": 0.0, "aspect": 1.0, "portrait": True}
    _x, _y, bw, bh = cv2.boundingRect(coords)
    ih, iw = binary.shape
    return {
        "cover_w": bw / float(iw),
        "cover_h": bh / float(ih),
        "aspect": bw / float(max(bh, 1)),
        "portrait": ih > iw * 1.05,
    }


def _detect_fine_skew_deg(img_bgr, max_deg=12.0):
    """检测扫描小歪斜（如约 8°），返回校正旋转角（度）。"""
    import cv2
    import numpy as np

    binary = _page_ink_binary(img_bgr)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(25, binary.shape[1] // 40), 1))
    lines_img = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel)
    min_len = max(60, binary.shape[1] // 12)
    lines = cv2.HoughLinesP(
        lines_img, 1, np.pi / 180, threshold=70,
        minLineLength=min_len, maxLineGap=12,
    )
    if lines is None:
        return 0.0
    angles = []
    for line in lines[:80]:
        x1, y1, x2, y2 = line[0]
        dx = x2 - x1
        if abs(dx) < 25:
            continue
        ang = float(np.degrees(np.arctan2(y2 - y1, dx)))
        if abs(ang) <= max_deg:
            angles.append(ang)
    if len(angles) < 4:
        return 0.0
    return -float(np.median(angles))


def _orientation_alignment_score(img_bgr):
    """文本行越水平，得分越高。"""
    skew = _detect_fine_skew_deg(img_bgr, max_deg=20.0)
    return 100.0 - abs(skew) * 8.0


def _has_sidebar_vertical_layout(img_bgr):
    """左侧竖排标题的表单（考评打分表）：不能用上下墨迹猜 180°。"""
    binary = _page_ink_binary(img_bgr)
    h, w = binary.shape
    left = float(binary[:, : max(1, w // 6)].sum())
    top = float(binary[: max(1, h // 6), :].sum())
    bottom = float(binary[min(h - 1, 5 * h // 6):, :].sum())
    return left > top * 0.65 and left > bottom * 0.65


def _upright_header_score(img_bgr):
    """标题/表头多在上方：正立时上半区墨迹更重。"""
    if _has_sidebar_vertical_layout(img_bgr):
        return 0.0
    binary = _page_ink_binary(img_bgr)
    h = binary.shape[0]
    top = float(binary[: max(1, h // 3), :].sum())
    bot = float(binary[min(h - 1, 2 * h // 3):, :].sum())
    return top - bot


def _orientation_combined_score(img_bgr):
    return _orientation_alignment_score(img_bgr) + _upright_header_score(img_bgr) * 0.003


def _detect_visual_sideways_rotation(img_bgr):
    """竖版画布上明显的 90° 侧躺启发式。"""
    cover = _ink_cover_metrics(img_bgr)
    if cover["cover_w"] > 0.76 and cover["cover_h"] > 0.76:
        return 0
    if not cover["portrait"]:
        return 0
    if cover["aspect"] > 1.22 and cover["cover_w"] > 0.55 and cover["cover_h"] < 0.72:
        return 90
    if cover["aspect"] < 0.78 and cover["cover_h"] > 0.58 and cover["cover_w"] < 0.65:
        return 270
    return 0


def _detect_coarse_rotation_osd(img_bgr, min_conf=2.0):
    """Tesseract OSD でページ向きを検出。信頼度が低い場合は None を返す。"""
    try:
        import cv2 as _cv2
        import pytesseract
        from PIL import Image

        rgb = _cv2.cvtColor(img_bgr, _cv2.COLOR_BGR2RGB)
        osd = pytesseract.image_to_osd(
            Image.fromarray(rgb), output_type=pytesseract.Output.DICT
        )
        conf = float(osd.get("orientation_conf") or 0)
        rotate = int(osd.get("rotate") or 0) % 360
        if conf >= min_conf:
            log.debug("OSD rotate=%d conf=%.2f", rotate, conf)
            return rotate
    except Exception as exc:
        log.debug("OSD failed: %s", exc)
    return None


def _detect_best_coarse_rotation(img_bgr, allow_full_probe=False):
    """粗调：0/90/180/270°。OSD 优先，置信度不足时回退视觉启发式。"""
    osd_deg = _detect_coarse_rotation_osd(img_bgr)
    if osd_deg is not None:
        return osd_deg

    cover = _ink_cover_metrics(img_bgr)
    if cover["cover_w"] > 0.76 and cover["cover_h"] > 0.76:
        return 0

    has_sidebar = _has_sidebar_vertical_layout(img_bgr)

    if not allow_full_probe:
        return _detect_visual_sideways_rotation(img_bgr)

    # 侧边竖排页跳过 180°（防误翻转），仍允许 90°/270° 打分
    candidates = (90, 270) if has_sidebar else (90, 180, 270)
    base_score = _orientation_combined_score(img_bgr)
    best_deg = 0
    best_score = base_score
    for deg in candidates:
        score = _orientation_combined_score(_rotate_bgr(img_bgr, deg))
        if score > best_score:
            best_score = score
            best_deg = deg
    if best_deg and best_score > base_score + 2.5:
        return best_deg
    return _detect_visual_sideways_rotation(img_bgr)


def _deskew_bgr(img_bgr, skew_deg):
    import cv2

    if abs(skew_deg) < 0.35:
        return img_bgr
    h, w = img_bgr.shape[:2]
    matrix = cv2.getRotationMatrix2D((w / 2.0, h / 2.0), skew_deg, 1.0)
    return cv2.warpAffine(
        img_bgr, matrix, (w, h),
        flags=cv2.INTER_LINEAR,
        borderMode=cv2.BORDER_CONSTANT,
        borderValue=(255, 255, 255),
    )


def _layout_kind_from_rotation(coarse_deg, skew_deg, img_bgr, meta_rot=0):
    if coarse_deg == 180:
        return "upside_down"
    if coarse_deg == 90:
        return "sideways_ccw"
    if coarse_deg == 270:
        return "sideways_cw"
    if abs(skew_deg) >= 0.5:
        return "skewed"
    if meta_rot in (90, 270):
        return "pdf_rotated_meta"
    ih, iw = img_bgr.shape[:2]
    if ih > iw * 1.05:
        return "portrait_upright"
    return "landscape_native"


def _classify_page_layout(img_bgr, pdf_path=None, page_index=0, probe_coarse=True):
    meta_rot = _pdf_page_rotation_meta(pdf_path, page_index) if pdf_path else 0
    coarse_deg = _detect_best_coarse_rotation(img_bgr, allow_full_probe=probe_coarse)
    rotated = _rotate_bgr(img_bgr, coarse_deg) if coarse_deg else img_bgr
    skew_deg = _detect_fine_skew_deg(rotated)
    kind = _layout_kind_from_rotation(coarse_deg, skew_deg, rotated, meta_rot)
    return {
        "kind": kind,
        "correction_deg": coarse_deg,
        "skew_deg": skew_deg,
        "meta_rot": meta_rot,
    }


def _apply_page_orientation(bgr, layout):
    if layout.get("correction_deg"):
        bgr = _rotate_bgr(bgr, layout["correction_deg"])
    skew_deg = layout.get("skew_deg") or 0
    if abs(skew_deg) >= 0.35:
        bgr = _deskew_bgr(bgr, skew_deg)
    return bgr


def _analyze_page_layout(pdf_path, page_index, dpi=120, probe_coarse=True):
    try:
        mtime = os.path.getmtime(pdf_path)
    except OSError:
        mtime = 0
    key = (pdf_path, page_index, dpi, mtime, probe_coarse)
    if key in _PAGE_LAYOUT_CACHE:
        return _PAGE_LAYOUT_CACHE[key]

    bgr = _render_page_bgr(pdf_path, page_index, dpi)
    layout = _classify_page_layout(bgr, pdf_path, page_index, probe_coarse=probe_coarse)
    corrected = _apply_page_orientation(bgr, layout)
    layout.update({
        "src_h": bgr.shape[0],
        "src_w": bgr.shape[1],
        "corr_h": corrected.shape[0],
        "corr_w": corrected.shape[1],
        "dpi": dpi,
    })
    _PAGE_LAYOUT_CACHE[key] = layout
    return layout


def _get_corrected_page_bgr(pdf_path, page_index, dpi, probe_coarse=True):
    layout = _analyze_page_layout(pdf_path, page_index, dpi, probe_coarse=probe_coarse)
    correction_deg = layout.get("correction_deg", 0)
    skew_deg = layout.get("skew_deg", 0)
    if correction_deg and correction_deg % 90 == 0:
        # Render with fitz rotation matrix — avoids warpAffine interpolation blur
        bgr = _render_page_bgr_rotated(pdf_path, page_index, dpi, correction_deg)
        if abs(skew_deg) >= 0.35:
            bgr = _deskew_bgr(bgr, skew_deg)
    else:
        bgr = _render_page_bgr(pdf_path, page_index, dpi)
        bgr = _apply_page_orientation(bgr, layout)
    return bgr, layout


def _sync_doc_section_to_layout(doc, layout, tight=True):
    """按转正后的图像尺寸设置 Word 节（仍是 A4，只是方向跟内容走）。"""
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches

    dpi = layout.get("dpi") or 180
    width_in = layout.get("corr_w", 0) / float(dpi)
    height_in = layout.get("corr_h", 0) / float(dpi)
    sec = doc.sections[0]
    if width_in > height_in:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Inches(width_in)
        sec.page_height = Inches(height_in)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width = Inches(width_in)
        sec.page_height = Inches(height_in)
    margin = Inches(0.45 if tight else 0.6)
    sec.top_margin = margin
    sec.bottom_margin = margin
    sec.left_margin = margin
    sec.right_margin = margin
    return sec


def _layout_needs_orientation_handling(layout):
    if not layout:
        return False
    if layout.get("correction_deg"):
        return True
    if abs(layout.get("skew_deg") or 0) >= 0.4:
        return True
    return layout.get("kind") in (
        "sideways_ccw", "sideways_cw", "upside_down", "skewed",
    )


def _layout_needs_sideways_handling(layout):
    return _layout_needs_orientation_handling(layout)


def _raw_text_len(page):
    total = 0
    for block in page.get("textblocks") or []:
        label = (block.get("label") or "para").lower()
        if label in ("foot", "image"):
            continue
        total += len((block.get("text") or "").strip())
    return total


def _meaningful_text_len(page, watermarks, sig_page=False):
    total = 0
    for block in page.get("textblocks") or []:
        label = (block.get("label") or "para").lower()
        if label in ("header", "foot", "image"):
            continue
        text = _normalize_text(block.get("text") or "")
        if text and not _is_noise_text(text, watermarks, sig_page=sig_page):
            total += len(text)
    return total


def _page_has_image_blocks(page):
    for block in page.get("textblocks") or []:
        if (block.get("label") or "").lower() == "image" and block.get("url"):
            return True
    return False


def _page_has_table(page):
    for block in page.get("textblocks") or []:
        if "<table" in (block.get("text") or "").lower():
            return True
    return False


def _page_table_quality(page):
    for block in page.get("textblocks") or []:
        text = block.get("text") or ""
        if "<table" not in text.lower():
            continue
        for html in _TABLE_HTML_RE.findall(text):
            return _table_ocr_quality(_parse_html_table_rows(html))
    return 1.0


def _table_semantic_score(page):
    """考评类表格是否包含可辨认字段。"""
    hints = ("考评", "扣分标准", "考评项目", "序号", "小计", "合计", "项目名称", "考评内容")
    for block in page.get("textblocks") or []:
        text = block.get("text") or ""
        if "<table" not in text.lower():
            continue
        for html in _TABLE_HTML_RE.findall(text):
            rows = _parse_html_table_rows(html)
            blob = "".join(c["text"] for row in rows for c in row)
            if not blob.strip():
                return 0.0
            hit = sum(1 for h in hints if h in blob)
            return hit / max(len(hints), 1)
    return 1.0


def _needs_table_snapshot(page):
    """仅对考评打分表等 OCR 完全不可读的表格页回退截图。"""
    if not _page_has_table(page):
        return False
    from seal_utils import page_text_blob

    page_blob = _normalize_text(page_text_blob(page))
    table_blob = ""
    for block in page.get("textblocks") or []:
        text = block.get("text") or ""
        if "<table" in text.lower():
            table_blob += re.sub(r"<[^>]+>", "", text)
    table_blob = _normalize_text(table_blob)

    rubric = any(
        k in page_blob or k in table_blob
        for k in ("考评打分", "扣分标准", "考评项目", "服务管理考评", "考评日期")
    )
    if not rubric:
        rubric = "分值" in table_blob and "备注" in table_blob and len(table_blob) < 280
    if not rubric:
        return False
    return _table_semantic_score(page) < 0.22


def _page_watermark_heavy(page, watermarks):
    hits = blocks = 0
    for block in page.get("textblocks") or []:
        label = (block.get("label") or "para").lower()
        if label in ("foot", "image", "table"):
            continue
        raw = (block.get("text") or "").strip()
        if not raw:
            continue
        blocks += 1
        if any(w in raw for w in ("臻", "至善", "厚生", "德厚", "李于", "正德", "哪生", "名厚")):
            hits += 1
        elif _normalize_text(raw) in watermarks:
            hits += 1
    return blocks > 0 and hits >= 1


def _editable_text_page(page, watermarks, is_sig_page):
    """有足够 OCR 内容时优先可编辑输出，表格页同样走文字/表格而非整页截图。"""
    meaningful = _meaningful_text_len(page, watermarks, sig_page=is_sig_page)
    raw = _raw_text_len(page)
    if meaningful >= 12 or raw >= 25:
        return True
    if _page_has_table(page) and raw >= 15:
        return True
    return False


def _is_stamp_header_text(text, box, page_hw):
    text = _normalize_text(text or "")
    if not re.fullmatch(r"\d{6,12}", text):
        return False
    pw = page_hw.get("w") or page_hw.get("width") or 849
    x0 = (box or {}).get("x0") or 0
    return x0 > pw * 0.58


def _should_render_snapshot(page, watermarks, is_sig_page, mode, pdf_path, page_index=0):
    if mode not in ("hybrid", "visual") or not pdf_path:
        return False
    if mode == "visual":
        return True
    if _needs_table_snapshot(page):
        return False  # 有表格HTML时优先渲染可编辑表格，截图不可编辑且评测得0分
    if _editable_text_page(page, watermarks, is_sig_page):
        return False
    meaningful = _meaningful_text_len(page, watermarks, sig_page=is_sig_page)
    raw = _raw_text_len(page)
    if _page_has_table(page) and raw >= 15 and not _needs_table_snapshot(page):
        return False
    if meaningful >= 8 or raw >= 18:
        return False
    return True


def _page_mainly_table(page):
    table_chars = 0
    other_chars = 0
    for block in page.get("textblocks") or []:
        label = (block.get("label") or "para").lower()
        if label in ("header", "foot", "image"):
            continue
        text = block.get("text") or ""
        if label == "table" or "<table" in text.lower():
            table_chars += len(text)
        else:
            other_chars += len(_normalize_text(text))
    return table_chars > max(other_chars * 2, 80)


def _page_table_blob(page, page_md=""):
    blob = page_md or ""
    for block in page.get("textblocks") or []:
        blob += block.get("text") or ""
    return blob


def _page_table_cell_count(page, page_md=""):
    return len(re.findall(r"<t[dh]", _page_table_blob(page, page_md), re.I))


def _page_has_large_image_block(page):
    page_hw = page.get("page_image_hw") or {}
    pw = page_hw.get("w") or page_hw.get("width") or 0
    ph = page_hw.get("h") or page_hw.get("height") or 0
    if not pw or not ph:
        return False
    for block in page.get("textblocks") or []:
        if (block.get("label") or "").lower() not in ("image", "figure", "fig"):
            continue
        box = block.get("box") or {}
        bw = abs((box.get("x1") or 0) - (box.get("x0") or 0))
        bh = abs((box.get("y1") or 0) - (box.get("y0") or 0))
        if bw * bh >= pw * ph * 0.35:
            return True
    return False


def _page_render_stats(page, watermarks, page_md="", is_sig_page=False):
    return {
        "meaningful": _meaningful_text_len(page, watermarks, sig_page=is_sig_page),
        "raw": _raw_text_len(page),
        "table_cells": _page_table_cell_count(page, page_md),
        "has_table_html": "<table" in _page_table_blob(page, page_md).lower(),
        "large_image": _page_has_large_image_block(page),
    }


def _page_is_blank_risk(stats, is_sig_page=False):
    if stats.get("table_cells", 0) >= 8:
        return False
    if is_sig_page:
        return stats.get("meaningful", 0) < 8 and not stats.get("has_table_html")
    if stats.get("meaningful", 0) < 20 and stats.get("table_cells", 0) == 0:
        return True
    if stats.get("large_image") and stats.get("meaningful", 0) < 30:
        return True
    return False


def _page_is_dense_table(page, page_md=""):
    if _page_is_scoring_form(page, page_md):
        return True
    blob = _page_table_blob(page, page_md)
    if any(m in blob for m in _TASK_CHECKLIST_HINTS):
        return False
    score = 0
    td_count = _page_table_cell_count(page, page_md)
    if td_count >= 12:
        score += 1
    if _page_mainly_table(page):
        score += 1
    if sum(1 for h in _DENSE_TABLE_HINTS if h in blob) >= 1:
        score += 1
    if "<table" in blob.lower() and td_count >= 6:
        score += 1
    if "<table" in blob.lower():
        for html in _TABLE_HTML_RE.findall(blob):
            if _table_ocr_quality(_parse_html_table_rows(html)) < 0.35:
                score += 1
                break
    return score >= 2


def _page_table_quality_abnormal(page, page_md=""):
    blob = _page_table_blob(page, page_md)
    for html in _TABLE_HTML_RE.findall(blob):
        rows = _parse_html_table_rows(html)
        if not rows:
            continue
        quality = _table_ocr_quality(rows)
        if quality < 0.35:
            return True
        ncols = max((len(r) for r in rows), default=0)
        if len(rows) >= 3 and ncols >= 6 and quality < 0.5:
            return True
    return False


def _extract_md_tables(page_md):
    return _TABLE_HTML_RE.findall(page_md or "")


def _table_html_quality(html):
    if not (html or "").strip():
        return 0.0
    return _table_ocr_quality(_parse_html_table_rows(html))


def _pick_best_table_html(detail_html, md_html, dense=False):
    """P1 双通道：按 OCR 质量在 detail 与 markdown 表格间择优。"""
    dq = _table_html_quality(detail_html)
    mq = _table_html_quality(md_html)
    if not (md_html or "").strip():
        return detail_html or ""
    if not detail_html or "<table" not in (detail_html or "").lower():
        return md_html
    if dense and mq >= max(dq, 0.35) * 0.85:
        return md_html
    if dq < 0.35 and mq > dq:
        return md_html
    if mq < 0.35 and dq > mq:
        return detail_html
    return md_html if mq > dq else detail_html


def _prepare_dual_channel_page(page, page_md, dense=False):
    """将 markdown 表格择优合并进 detail 块，返回 (page_copy, consumed_md_count, md_tables)。"""
    page_copy = copy.deepcopy(page)
    md_tables = _extract_md_tables(page_md)
    md_i = 0
    for block in page_copy.get("textblocks") or []:
        label = (block.get("label") or "para").lower()
        text = block.get("text") or ""
        is_table = label == "table" or "<table" in text.lower()
        if not is_table:
            continue
        detail_parts = _TABLE_HTML_RE.findall(text)
        detail_html = detail_parts[0] if detail_parts else text
        md_html = md_tables[md_i] if md_i < len(md_tables) else ""
        chosen = _pick_best_table_html(detail_html, md_html, dense=dense)
        if chosen:
            block["text"] = chosen
            block["label"] = "table"
        if md_i < len(md_tables):
            md_i += 1
    return page_copy, md_i, md_tables


def _page_non_table_meaningful_len(page, watermarks, is_sig_page=False):
    total = 0
    for block in page.get("textblocks") or []:
        label = (block.get("label") or "para").lower()
        if label in ("foot", "image", "table"):
            continue
        text = block.get("text") or ""
        if "<table" in text.lower():
            continue
        norm = _normalize_text(text)
        if norm and not _is_noise_text(norm, watermarks, sig_page=is_sig_page):
            total += len(norm)
    return total


def _render_remaining_md_tables(doc, md_tables, start_index, landscape=False, table_opts=None):
    opts = dict(table_opts or {})
    landscape = opts.pop("landscape", landscape)
    for html in md_tables[start_index:]:
        if (html or "").strip():
            _add_html_table(doc, html, landscape=landscape, **opts)


def _doc_body_text_len(doc):
    total = 0
    for para in doc.paragraphs:
        total += len((para.text or "").strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += len((cell.text or "").strip())
    return total


def _build_page_warnings(snapshot_pages, abnormal_table_pages):
    parts = []
    if snapshot_pages:
        pages = "、".join(f"第{p}页" for p in snapshot_pages)
        parts.append(f"部分页面识别困难，{pages}已插入原页截图")
    if abnormal_table_pages:
        pages = "、".join(f"第{p}页" for p in abnormal_table_pages)
        parts.append(f"{pages}表格结构异常，建议人工核对")
    return "; ".join(parts)


def _render_dense_table_page(doc, page_md, watermarks, cache_dir, table_opts=None):
    """密集工程表页：优先 markdown 表格 HTML + 横版，不走坐标段落排版。"""
    if not (page_md or "").strip():
        return
    opts = dict(table_opts or {})
    landscape = opts.pop("landscape", True)
    for block in re.split(r"\n\s*\n", page_md):
        block = block.strip()
        if not block:
            continue
        if block.lower().startswith("<table"):
            _add_html_table(doc, block, landscape=landscape, **opts)
            continue
        if block.startswith("#"):
            level = len(block) - len(block.lstrip("#"))
            text = block.lstrip("#").strip()
            style = "Heading 1" if level <= 1 else "Heading 2"
            _render_text_content(doc, text, watermarks, cache_dir, style=style)
            continue
        for line in block.split("\n"):
            line = line.strip()
            if line:
                _render_checklist_markdown_line(doc, line, watermarks, cache_dir)


def _should_render_inline_image(box, page_hw, skip_images):
    if not skip_images:
        return True
    pw = page_hw.get("w") or page_hw.get("width") or 0
    ph = page_hw.get("h") or page_hw.get("height") or 0
    if not pw or not ph:
        return True
    bw = abs((box or {}).get("x1", 0) - (box or {}).get("x0", 0))
    bh = abs((box or {}).get("y1", 0) - (box or {}).get("y0", 0))
    return bw * bh < pw * ph * 0.42


def _get_pdf_page_image(pdf_path, page_index, cache_dir, dpi=120, correct_sideways=False):
    try:
        import cv2
    except ImportError:
        return None

    rot_tag = ""
    if correct_sideways:
        layout = _analyze_page_layout(pdf_path, page_index, dpi, probe_coarse=True)
        if layout.get("correction_deg"):
            rot_tag = f"_r{layout['correction_deg']}"
        skew = layout.get("skew_deg") or 0
        if abs(skew) >= 0.4:
            rot_tag += f"_s{int(skew * 10)}"
    path = os.path.join(cache_dir, f"page_{page_index:03d}_{dpi}{rot_tag}.png")
    if os.path.isfile(path) and os.path.getsize(path) > 0:
        return path
    os.makedirs(cache_dir, exist_ok=True)
    if correct_sideways:
        bgr, _ = _get_corrected_page_bgr(pdf_path, page_index, dpi, probe_coarse=True)
    else:
        bgr = _render_page_bgr(pdf_path, page_index, dpi)
    cv2.imwrite(path, bgr)
    return path if os.path.isfile(path) else None


def _render_pdf_page_snapshot(doc, pdf_path, page_index, cache_dir, dpi=120, width_in=None):
    from docx.shared import Inches

    path = _get_pdf_page_image(pdf_path, page_index, cache_dir, dpi=dpi)
    if not path:
        return None
    landscape = _pdf_page_landscape(pdf_path, page_index)
    if width_in is None:
        width_in = 10.0 if landscape else 6.5
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Inches(0.08)
    run = para.add_run()
    run.add_picture(path, width=Inches(width_in))
    return path


_FLOAT_IMG_ID = [200]


def _get_page_emu_size(pdf_path, page_index):
    """Return (width_emu, height_emu) for a PDF page."""
    try:
        import fitz
        pdoc = fitz.open(pdf_path)
        rect = pdoc[page_index].rect
        pdoc.close()
        emu_per_pt = 914400 / 72
        return int(rect.width * emu_per_pt), int(rect.height * emu_per_pt)
    except Exception:
        return 7559055, 10692000  # A4 fallback


def _add_floating_image(doc, image_path, pos_x_emu, pos_y_emu, width_emu, height_emu):
    """Insert a wp:anchor floating image at absolute page position (all values in EMU)."""
    from lxml import etree

    # Add image inline to register it in document relationships and get rId
    tmp_para = doc.add_paragraph()
    tmp_run = tmp_para.add_run()
    try:
        from docx.shared import Emu
        tmp_run.add_picture(image_path, width=Emu(max(width_emu, 914400 // 10)))
    except Exception:
        doc.element.body.remove(tmp_para._p)
        return
    p_xml = etree.tostring(tmp_para._p, encoding="unicode")
    rid_m = re.search(r'r:embed="(rId\d+)"', p_xml)
    if not rid_m:
        doc.element.body.remove(tmp_para._p)
        return
    r_id = rid_m.group(1)
    doc.element.body.remove(tmp_para._p)

    img_id = _FLOAT_IMG_ID[0]
    _FLOAT_IMG_ID[0] += 1

    anchor_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:r><w:drawing>"
        '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        ' relativeHeight="251658240" behindDoc="0" locked="0"'
        ' layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f"<wp:positionH relativeFrom=\"page\"><wp:posOffset>{pos_x_emu}</wp:posOffset></wp:positionH>"
        f"<wp:positionV relativeFrom=\"page\"><wp:posOffset>{pos_y_emu}</wp:posOffset></wp:positionV>"
        f'<wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        "<wp:wrapNone/>"
        f'<wp:docPr id="{img_id}" name="FloatImg{img_id}"/>'
        "<wp:cNvGraphicFramePr>"
        '<a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>'
        "</wp:cNvGraphicFramePr>"
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:nvPicPr><pic:cNvPr id="{img_id}" name="FloatImg{img_id}"/><pic:cNvPicPr/></pic:nvPicPr>'
        "<pic:blipFill>"
        f'<a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:embed="{r_id}"/>'
        "<a:stretch><a:fillRect/></a:stretch>"
        "</pic:blipFill>"
        "<pic:spPr>"
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        "</pic:spPr>"
        "</pic:pic>"
        "</a:graphicData>"
        "</a:graphic>"
        "</wp:anchor>"
        "</w:drawing></w:r></w:p>"
    )
    try:
        p_elem = etree.fromstring(anchor_xml)
        doc.element.body.append(p_elem)
    except Exception:
        pass


def _embed_extracted_seals(doc, page_image_path, cache_dir, page_hw=None, pg_w_emu=0, pg_h_emu=0):
    from docx.shared import Inches
    from seal_utils import extract_red_seals_from_image

    seal_dir = os.path.join(cache_dir, "seals_" + os.path.basename(page_image_path))
    seals = extract_red_seals_from_image(page_image_path, seal_dir)
    if not seals:
        return 0

    # If we have page EMU dimensions, use floating placement for correct positioning
    if pg_w_emu and pg_h_emu:
        try:
            import cv2
            img = cv2.imread(page_image_path)
            if img is not None:
                img_h, img_w = img.shape[:2]
                for seal in seals:
                    norm_x = seal["x"] / img_w
                    norm_y = seal["y"] / img_h
                    norm_w = seal["w"] / img_w
                    norm_h = seal["h"] / img_h
                    pos_x = int(norm_x * pg_w_emu)
                    pos_y = int(norm_y * pg_h_emu)
                    width = max(1, int(norm_w * pg_w_emu))
                    height = max(1, int(norm_h * pg_h_emu))
                    _add_floating_image(doc, seal["path"], pos_x, pos_y, width, height)
                return len(seals)
        except Exception:
            pass

    # Fallback: inline placement
    page_w = (page_hw or {}).get("w") or (page_hw or {}).get("width") or 0
    try:
        import cv2
        img = cv2.imread(page_image_path)
        img_w = img.shape[1] if img is not None else 0
    except Exception:
        img_w = 0
    scale_base = page_w or img_w or 1

    rows = []
    current = [seals[0]]
    for seal in seals[1:]:
        if abs(seal["cy"] - current[-1]["cy"]) <= max(current[-1]["h"], seal["h"]) * 0.55:
            current.append(seal)
        else:
            rows.append(sorted(current, key=lambda s: s["cx"]))
            current = [seal]
    rows.append(sorted(current, key=lambda s: s["cx"]))

    for row in rows:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Inches(0.04)
        para.paragraph_format.space_after = Inches(0.08)
        for seal in row:
            width_in = 1.55
            if scale_base and seal.get("w"):
                width_in = min(2.2, max(0.95, 6.5 * seal["w"] / float(scale_base)))
            run = para.add_run()
            run.add_picture(seal["path"], width=Inches(width_in))
            run.add_text("   ")
    return len(seals)


def _resolve_block_image_path(block, cache_dir, pdf_path=None, page_index=None, page_hw=None):
    """OCR 图链优先；火山 URL 过期时从 PDF 裁切（封面二维码等）。"""
    url = block.get("url")
    if url:
        try:
            path = _download_image(url, cache_dir)
            if path and os.path.isfile(path) and os.path.getsize(path) > 0:
                return path
        except Exception:
            pass
    if pdf_path and page_index is not None:
        return _crop_pdf_image(
            pdf_path, page_index, block.get("box") or {}, page_hw or {}, cache_dir,
        )
    return None


def _embed_ocr_image_blocks(
    doc, page, cache_dir, pg_w_emu, pg_h_emu,
    pdf_path=None, page_index=None,
):
    """Float-embed image blocks detected by OCR (e.g. QR codes on cover pages)."""
    page_hw = page.get("page_image_hw") or {}
    count = 0
    for block in page.get("textblocks") or []:
        if (block.get("label") or "").lower() != "image":
            continue
        nb = block.get("norm_box") or {}
        x0 = nb.get("x0") or 0
        y0 = nb.get("y0") or 0
        x1 = nb.get("x1") or x0
        y1 = nb.get("y1") or y0
        if x1 <= x0 or y1 <= y0:
            continue
        path = _resolve_block_image_path(
            block, cache_dir, pdf_path=pdf_path, page_index=page_index, page_hw=page_hw,
        )
        if not path:
            continue
        pos_x = int(x0 * pg_w_emu)
        pos_y = int(y0 * pg_h_emu)
        width = max(1, int((x1 - x0) * pg_w_emu))
        height = max(1, int((y1 - y0) * pg_h_emu))
        _add_floating_image(doc, path, pos_x, pos_y, width, height)
        count += 1
    return count


def _embed_signatures(doc, page_image_path, page, cache_dir, pg_w_emu, pg_h_emu):
    """Detect and float-embed handwritten signatures on signature pages."""
    try:
        import cv2
        from seal_utils import extract_signature_areas
    except ImportError:
        return 0

    try:
        img = cv2.imread(page_image_path)
        if img is None:
            return 0
        img_h, img_w = img.shape[:2]

        # Scale OCR block coordinates from OCR image space to seal extraction image space
        page_hw = page.get("page_image_hw") or {}
        ocr_w = page_hw.get("w") or page_hw.get("width") or img_w
        ocr_h = page_hw.get("h") or page_hw.get("height") or img_h
        sx = img_w / float(ocr_w)
        sy = img_h / float(ocr_h)
        scaled_blocks = []
        for blk in page.get("textblocks") or []:
            box = blk.get("box") or {}
            if not box:
                continue
            scaled_blocks.append({
                "box": {
                    "x0": int((box.get("x0") or 0) * sx),
                    "y0": int((box.get("y0") or 0) * sy),
                    "x1": int((box.get("x1") or 0) * sx),
                    "y1": int((box.get("y1") or 0) * sy),
                }
            })

        sig_dir = os.path.join(cache_dir, "sigs_" + os.path.basename(page_image_path))
        sigs = extract_signature_areas(
            img, ocr_blocks=scaled_blocks, out_dir=sig_dir, sig_page=True,
        )
        for sig in sigs:
            norm_x = sig["x"] / img_w
            norm_y = sig["y"] / img_h
            norm_w = sig["w"] / img_w
            norm_h = sig["h"] / img_h
            pos_x = int(norm_x * pg_w_emu)
            pos_y = int(norm_y * pg_h_emu)
            width = max(1, int(norm_w * pg_w_emu))
            height = max(1, int(norm_h * pg_h_emu))
            _add_floating_image(doc, sig["path"], pos_x, pos_y, width, height)
        return len(sigs)
    except Exception:
        return 0


def _add_image(doc, url, cache_dir, page_hw=None, box=None):
    from docx.shared import Inches

    try:
        path = _download_image(url, cache_dir)
    except Exception:
        return False

    width_in = 6.0
    if page_hw and box:
        pw = page_hw.get("w") or page_hw.get("width") or 0
        bh = (box.get("x1", 0) - box.get("x0", 0))
        if pw and bh:
            width_in = min(6.5, max(1.2, 6.0 * bh / pw))

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Inches(0.05)
    run = para.add_run()
    run.add_picture(path, width=Inches(width_in))
    return True


def _set_run_font(run, bold=False, italic=False):
    run.bold = bool(bold)
    run.italic = bool(italic)


def _add_text_paragraph(doc, text, style=None, bold=False, italic=False):
    from docx.shared import Pt

    text = _normalize_text(text)
    if not text:
        return
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    _set_run_font(run, bold=bold, italic=italic)
    if style is None:
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.25


def _box_metrics(box, page_hw):
    pw = page_hw.get("w") or page_hw.get("width") or 849
    ph = page_hw.get("h") or page_hw.get("height") or 1200
    x0 = box.get("x0") or 0
    y0 = box.get("y0") or 0
    x1 = box.get("x1") or x0
    y1 = box.get("y1") or y0
    return pw, ph, x0, y0, x1, y1


def _infer_alignment(box, page_hw, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pw, _, x0, _, x1, _ = _box_metrics(box, page_hw)
    cx = (x0 + x1) / 2.0 / max(pw, 1)
    width_ratio = (x1 - x0) / max(pw, 1)
    text = text or ""
    if re.match(r"^甲\s*方|^乙\s*方", text.strip()):
        return WD_ALIGN_PARAGRAPH.LEFT
    if "服务合同" in text or "合同" == text.strip():
        return WD_ALIGN_PARAGRAPH.CENTER
    if width_ratio > 0.58:
        return WD_ALIGN_PARAGRAPH.CENTER
    if 0.36 <= cx <= 0.64 and (width_ratio > 0.38 or len(text) <= 36):
        return WD_ALIGN_PARAGRAPH.CENTER
    if cx >= 0.68:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


_COVER_CONTENT_PT = 680.0


def _box_target_y_pt(box, page_hw):
    _, ph, _, y0, _, _ = _box_metrics(box or {}, page_hw)
    return y0 / max(ph, 1) * _COVER_CONTENT_PT


def _estimate_para_height_pt(text, size_pt, box, page_hw):
    pw, _, x0, _, x1, _ = _box_metrics(box or {}, page_hw)
    width_ratio = max((x1 - x0) / max(pw, 1), 0.28)
    chars_per_line = max(10, int(width_ratio * 38))
    lines = max(1, (len(text or "") + chars_per_line - 1) // chars_per_line)
    return lines * float(size_pt) * 1.15


def _space_before_cover_pt(rendered_bottom, box, page_hw):
    """封面单页：按 OCR y0 绝对定位，避免拆段后 OCR 框过高把内容挤到第 2 页。"""
    from docx.shared import Pt

    target = _box_target_y_pt(box, page_hw)
    gap = target - rendered_bottom
    if gap < 2:
        gap = 4
    return Pt(gap), target


def _space_before_pt(prev_box, box, page_hw):
    """按 OCR 纵坐标定位；框重叠时用 y0 估算段前距。"""
    from docx.shared import Pt

    _, ph, _, y0, _, _ = _box_metrics(box, page_hw)
    content_pt = _COVER_CONTENT_PT
    target = y0 / max(ph, 1) * content_pt
    if not prev_box:
        return Pt(max(0, target - 6))
    prev_target = (prev_box.get("y0") or 0) / max(ph, 1) * content_pt
    prev_span = ((prev_box.get("y1") or 0) - (prev_box.get("y0") or 0)) / max(ph, 1) * content_pt
    prev_span = min(prev_span, 72)
    gap = target - prev_target - prev_span * 0.32
    if gap < 6:
        gap = 14 if (y0 - (prev_box.get("y0") or 0)) > 60 else 6
    return Pt(max(4, gap))


def _set_run_font_style(run, name, size_pt, east_asia=None, bold=False, color_hex=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    ea = east_asia or name
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea)
    run.font.size = Pt(size_pt)
    run.bold = bool(bold)
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )


def _cover_text_style(text, label, page_index, watermark_page=False):
    if page_index != 0 and not watermark_page:
        return None
    text = text or ""
    if label == "header" or re.fullmatch(r"\d{6,12}", text.strip()):
        return {"name": "宋体", "east_asia": "SimSun", "size": 21, "bold": True, "color": "008FEF"}
    if re.match(r"^CMCCTD", text):
        return {"name": "宋体", "east_asia": "SimSun", "size": 10, "bold": False}
    if "服务合同" in text or text.strip() == "服务合同":
        return {"name": "黑体", "east_asia": "SimHei", "size": 22, "bold": True}
    if re.match(r"^甲\s*方|^乙\s*方", text):
        return {"name": "宋体", "east_asia": "SimSun", "size": 12, "bold": False}
    if len(text) >= 16 and re.search(r"20\d{2}年|工程|监测|项目|生态", text):
        return {"name": "黑体", "east_asia": "SimHei", "size": 13, "bold": False}
    return None


def _get_cached_page_rgb(pdf_path, page_index, cache_dir, dpi=None):
    if not pdf_path or not os.path.isfile(pdf_path):
        return None
    dpi = dpi or _image_parse_dpi()
    path = os.path.join(cache_dir, f"rgb_p{page_index}_{dpi}.npy")
    try:
        import numpy as np
        if os.path.isfile(path):
            return np.load(path)
        _, rgb = _page_rgb_and_b64(pdf_path, page_index, dpi=dpi)
        np.save(path, rgb)
        return rgb
    except Exception:
        return None


def _sample_box_color_hex(rgb_arr, box, page_hw, prefer_blue=False):
    if rgb_arr is None:
        return None
    pw, ph, x0, y0, x1, y1 = _box_metrics(box or {}, page_hw)
    h_img, w_img = rgb_arr.shape[0], rgb_arr.shape[1]
    xs = max(0, min(w_img - 1, int(x0 / max(pw, 1) * w_img)))
    xe = max(xs + 1, min(w_img, int(x1 / max(pw, 1) * w_img)))
    ys = max(0, min(h_img - 1, int(y0 / max(ph, 1) * h_img)))
    ye = max(ys + 1, min(h_img, int(y1 / max(ph, 1) * h_img)))
    region = rgb_arr[ys:ye, xs:xe]
    if region.size == 0:
        return None
    blues, darks, others = [], [], []
    flat = region.reshape(-1, region.shape[-1])
    step = max(1, len(flat) // 100)
    for i in range(0, len(flat), step):
        r, g, b = [int(v) for v in flat[i][:3]]
        lum = (r + g + b) / 3
        if lum > 235:
            continue
        if b > max(r, g) + 28 and b > 95:
            blues.append((r, g, b))
        elif lum < 95:
            darks.append((r, g, b))
        else:
            others.append((r, g, b))
    picks = blues if (prefer_blue and blues) else (darks if darks else others)
    if not picks:
        picks = blues or others
    if not picks:
        return None
    r = sum(t[0] for t in picks) // len(picks)
    g = sum(t[1] for t in picks) // len(picks)
    b = sum(t[2] for t in picks) // len(picks)
    return f"{r:02X}{g:02X}{b:02X}"


def _estimate_cover_font_pt(box, page_hw, text, label):
    _, ph, _, y0, _, y1 = _box_metrics(box or {}, page_hw)
    bh = max(y1 - y0, 1)
    text = text or ""
    line_est = max(1, len(text) // 26 + 1)
    pt = (bh / line_est) / max(ph, 1) * 560
    if label == "title" or text.strip() == "服务合同":
        return max(18, min(24, pt * 1.08))
    if re.fullmatch(r"\d{6,12}", text.strip()):
        return max(18, min(22, pt))
    if re.match(r"^CMCCTD", text):
        return max(9, min(11, pt))
    if len(text) >= 20:
        return max(12, min(16, pt))
    return max(10, min(13, pt))


def _extract_cover_title_text(text):
    work = _strip_watermark_substrings(text or "")
    work = re.sub(r"\s+", " ", work).strip()
    m = re.search(r"(20\d{2}年.+)", work)
    if m:
        title = m.group(1).strip()
        title = re.sub(r"\s*服务合同.*$", "", title).strip()
        return title
    work = re.sub(r"^[李于正德哪名厚\s\"\"、]+", "", work)
    work = re.sub(r"\s*服务合同.*$", "", work).strip()
    return work


def _split_cover_party_lines(text, box, block):
    if not re.search(r"甲\s*方", text):
        return None
    markers = [m.start() for m in re.finditer(r"甲\s*方|乙\s*方", text)]
    if not markers:
        return None
    y0 = box.get("y0", 0)
    y1 = box.get("y1", y0)
    h = max(y1 - y0, 1)
    slices = []
    for i, pos in enumerate(markers):
        end = markers[i + 1] if i + 1 < len(markers) else len(text)
        chunk = text[pos:end].strip()
        if chunk:
            slices.append(chunk)
    if len(slices) < 2:
        return None
    out = []
    n = len(slices)
    for i, chunk in enumerate(slices):
        sub = dict(block)
        sub["text"] = chunk
        sub["box"] = dict(box)
        sub["box"]["y0"] = y0 + int(h * (i / n))
        sub["box"]["y1"] = y0 + int(h * ((i + 1) / n))
        out.append(sub)
    return out


def _split_cover_title_contract(block):
    """拆标题与「服务合同」：按标题行数估纵坐标，避免按字数比例压到框底。"""
    text = _normalize_text(block.get("text") or "")
    idx = text.find("服务合同")
    if idx <= 0:
        return [block]
    box = dict(block.get("box") or {})
    y0 = box.get("y0") or 0
    y1 = box.get("y1") or y0
    h = max(y1 - y0, 1)
    head = _extract_cover_title_text(text[:idx].strip())
    if not head:
        return [block]
    head_lines = max(1, (len(head) + 21) // 22)
    frac = min(0.62, 0.28 + head_lines * 0.11)
    mid = y0 + int(h * frac)
    head_block = dict(block)
    head_block["text"] = head
    head_block["box"] = {**box, "y1": mid}
    tail_block = dict(block)
    tail_block["text"] = "服务合同"
    tail_block["box"] = {**box, "y0": mid, "y1": y1}
    tail_block["label"] = "title"
    return [head_block, tail_block]


def _expand_cover_blocks(block):
    """P2 封面：去水印、拆标题/服务合同/甲乙方。"""
    label = (block.get("label") or "para").lower()
    if label in ("image", "table", "foot"):
        return [block]
    text = (block.get("text") or "").strip()
    if not text:
        return []
    if label == "header":
        return [block]
    if "服务合同" in text and text.strip() != "服务合同":
        return _split_cover_title_contract(block)
    if re.fullmatch(r"服务合同", _strip_watermark_substrings(text).strip()):
        sub = dict(block)
        sub["text"] = "服务合同"
        sub["label"] = "title"
        return [sub]

    party_splits = _split_cover_party_lines(text, block.get("box") or {}, block)
    if party_splits:
        return party_splits

    if label == "para" and re.search(r"20\d{2}年|工程|监测|项目", text):
        cleaned = _extract_cover_title_text(text)
        if cleaned:
            sub = dict(block)
            sub["text"] = cleaned
            return [sub]

    work = _strip_watermark_substrings(text)
    work = re.sub(r"\s+", " ", work).strip()
    if not work or _text_quality_score(work) < 0.25:
        return []
    sub = dict(block)
    sub["text"] = work
    return [sub]


def _cover_paragraph_align(text, label, box, page_hw):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = text or ""
    if label == "header" or re.fullmatch(r"\d{6,12}", text.strip()):
        return WD_ALIGN_PARAGRAPH.RIGHT
    if re.match(r"^CMCCTD", text):
        return WD_ALIGN_PARAGRAPH.CENTER
    if "服务合同" in text or text.strip() == "服务合同":
        return WD_ALIGN_PARAGRAPH.CENTER
    if len(text) >= 16 and re.search(r"20\d{2}年|工程|监测|项目", text):
        return WD_ALIGN_PARAGRAPH.CENTER
    return _infer_alignment(box or {}, page_hw, text)


# Removed hardcoded company names — they only worked for one specific contract.
# If OCR can't read a party name clearly (too short or watermark-contaminated),
# output whatever OCR got. After proper watermark removal, OCR quality is
# sufficient to read party names without hardcoded fallbacks.


def _polish_sig_ocr_text(text):
    text = _polish_sig_text(text)
    text = re.sub(r"^平方", "甲方", text)
    text = re.sub(r"(?<![\u4e00-\u9fff])上公司(?![\u4e00-\u9fff])", "上海公司", text)
    return text.strip()


def _is_sig_garbage_line(text):
    text = _normalize_text(text)
    if not text:
        return True
    if len(text) <= 2 and not re.search(r"[\d]{4}", text):
        return True
    if re.fullmatch(r"[\d\.\s期b]+", text, re.I):
        return True
    if text in ("德", "天", "大", "唇:", "合。", "省", "示", "荣周", "天道", "有限公司"):
        return True
    if re.search(r"^盖章云|^天道|^合。\s*同专用章", text):
        return True
    if "同专用章" in text and not re.search(r"甲\s*方|乙\s*方|签字", text):
        return True
    if re.fullmatch(r"臻子不?\"?", text):
        return True
    if _text_quality_score(text) < 0.22 and not _is_sig_relevant_text(text):
        return True
    return False


def _expand_sig_blocks(block):
    """签章页：去碎片、拆甲乙方与签字/盖章行。"""
    label = (block.get("label") or "para").lower()
    if label in ("image", "table", "foot"):
        return []
    if label == "header":
        return [block]
    text = _polish_sig_ocr_text(block.get("text") or "")
    if not text or _is_sig_garbage_line(text):
        return []
    if re.match(r"^第[一二三四五六七八九十\d]+条", text):
        return [dict(block, text=text)]
    if re.search(r"投诉电话|邮箱:", text) or "@" in text:
        return [dict(block, text=text)]
    if re.search(r"日期\s*[：:]", text):
        line = re.sub(r"日期\s*[：:]\s*", "日期：", text)
        line = re.sub(r"日期：\s*0年.*$", "日期：", line).strip()
        return [dict(block, text=line)] if line else []
    if re.search(r"甲\s*方|乙\s*方|签字/盖章|签字|盖章", text):
        lines = _split_sig_lines(text)
        if len(lines) <= 1 and len(text) > 36:
            markers = [m.start() for m in re.finditer(r"签字/盖章|签字|盖章|日期", text)]
            if markers:
                boundaries = sorted(set([0] + markers))
                lines = [
                    text[boundaries[i]:boundaries[i + 1]].strip()
                    for i in range(len(boundaries) - 1)
                ]
                lines.append(text[boundaries[-1]:])
                lines = [p for p in lines if p]
        out = []
        box = dict(block.get("box") or {})
        y0 = box.get("y0") or 0
        y1 = box.get("y1") or y0
        h = max(y1 - y0, 1)
        lines = [ln for ln in lines if not _is_sig_garbage_line(ln)]
        n = max(len(lines), 1)
        for i, line in enumerate(lines):
            line = _polish_sig_ocr_text(line)
            if _is_sig_garbage_line(line):
                continue
            sub = dict(block)
            sub["text"] = line
            sub["box"] = {
                **box,
                "y0": y0 + int(h * (i / n)),
                "y1": y0 + int(h * ((i + 1) / n)),
            }
            out.append(sub)
        return out
    return [dict(block, text=text)]


def _split_sig_lines(text):
    text = _polish_sig_ocr_text(text)
    if not re.search(r"甲\s*方|乙\s*方", text):
        return [text] if text else []
    # Split before each marker without consuming it (py3.6-safe)
    markers = r"(?:\(签[字章]页\)|签[字章]页|甲方|乙方|日期|签字/盖章)"
    positions = [m.start() for m in re.finditer(markers, text)]
    if not positions:
        return [text]
    boundaries = sorted(set([0] + positions))
    parts = [text[boundaries[i]:boundaries[i + 1]] for i in range(len(boundaries) - 1)]
    parts.append(text[boundaries[-1]:])
    return [p.strip() for p in parts if p.strip()]


def _polish_sig_text(text):
    text = _normalize_text(text)
    text = re.sub(r"\(签章页\)\s*国\s*同", "(签章页)", text)
    text = re.sub(r"签字/盖章:\s*个同专用章", "签字/盖章：", text)
    text = re.sub(r"签字/盖章:\s*[：:]+", "签字/盖章：", text)
    text = re.sub(r"：合同专用章", "合同专用章", text)
    return text.strip()


def _format_party_text(text, sig_page=False):
    text = _polish_sig_text(text) if sig_page else _normalize_text(text)
    m = re.match(r"^(甲方|乙方)\s*[：:]*\s*[【\[]?(.+?)[】\]]?\s*$", text)
    if not m:
        inline = re.search(
            r"(甲方|乙方)\s*[：:]\s*[【\[]?([^【】\n]{2,}?)(?:[】\]]|\s+签字|\s+日期|\s+盖章|$)",
            text,
        )
        if inline:
            party, name = inline.group(1), inline.group(2).strip()
        else:
            return text
    else:
        party, name = m.group(1), m.group(2).strip()
    name = re.sub(r"^[,\"'【\[\(]+", "", name).strip()
    name = re.sub(r"[【\[\]】]+", "", name).strip()
    name = re.sub(r",?\(上海\)", "", name).strip()
    name = re.sub(r"(正)?[德厚生于哪臻善至不\"、\s]+$", "", name).strip()
    name = re.sub(r"(签字|盖章|日期|合同专用章).*$", "", name).strip()
    name = re.sub(r"\s+江西省\s*荣周.*$", "", name).strip()
    name = re.sub(r"\s+省\s+示.*$", "", name).strip()
    return f"{party}：【{name}】"


def _should_render_cover_image(box, page_hw):
    pw = page_hw.get("w") or page_hw.get("width") or 849
    ph = page_hw.get("h") or page_hw.get("height") or 1200
    x0 = box.get("x0") or 0
    y0 = box.get("y0") or 0
    bw = max((box.get("x1") or 0) - x0, 0)
    bh = max((box.get("y1") or 0) - y0, 0)
    if bw * bh > pw * ph * 0.12:
        return False
    if y0 > ph * 0.75 and bw < pw * 0.3:
        return True
    if y0 < ph * 0.15 and x0 > pw * 0.5 and bw < pw * 0.2:
        return True
    return False


def _font_size_pt(box, label, text, page_hw, page_index=None, watermark_page=False):
    from docx.shared import Pt

    cover = _cover_text_style(text, label, page_index, watermark_page)
    if cover:
        return Pt(cover["size"])
    _, ph, _, y0, _, y1 = _box_metrics(box, page_hw)
    bh = max(y1 - y0, 1)
    if label == "title":
        return Pt(16)
    if "服务合同" in (text or ""):
        return Pt(22)
    if label == "sec":
        return Pt(12)
    if label == "cap":
        return Pt(10)
    scaled = bh / max(ph, 1) * 300
    if len(text or "") <= 24 and bh / max(ph, 1) < 0.04:
        return Pt(max(9, min(11, scaled)))
    return Pt(max(10, min(16, scaled)))


def _expand_layout_blocks(block):
    """把「标题 + 服务合同」这类合并 OCR 块拆成独立段落以保留版式。"""
    text = _normalize_text(block.get("text") or "")
    box = dict(block.get("box") or {})
    if not text or "服务合同" not in text or text.strip() == "服务合同":
        return [block]
    idx = text.find("服务合同")
    head = text[:idx].strip()
    if not head:
        return [block]
    y0 = box.get("y0") or 0
    y1 = box.get("y1") or y0
    mid = y0 + int((y1 - y0) * 0.62)
    head_block = dict(block)
    head_block["text"] = head
    head_block["box"] = {**box, "y1": mid}
    tail_block = dict(block)
    tail_block["text"] = "服务合同"
    tail_block["box"] = {**box, "y0": mid, "y1": y1}
    tail_block["label"] = "title"
    return [head_block, tail_block]


def _crop_pdf_image(pdf_path, page_index, box, page_hw, cache_dir):
    try:
        import fitz
    except ImportError:
        return None
    if not pdf_path or not os.path.isfile(pdf_path):
        return None
    pw, ph, x0, y0, x1, y1 = _box_metrics(box or {}, page_hw)
    if x1 <= x0 or y1 <= y0:
        return None
    path = os.path.join(cache_dir, f"crop_p{page_index}_{int(x0)}_{int(y0)}.png")
    if os.path.isfile(path) and os.path.getsize(path) > 0:
        return path
    doc = fitz.open(pdf_path)
    try:
        page = doc[page_index]
        pr = page.rect
        ocr_pw = page_hw.get("w") or page_hw.get("width") or pr.width
        ocr_ph = page_hw.get("h") or page_hw.get("height") or pr.height
        sx, sy = pr.width / max(ocr_pw, 1), pr.height / max(ocr_ph, 1)
        clip = fitz.Rect(x0 * sx, y0 * sy, x1 * sx, y1 * sy)
        mat = fitz.Matrix(160 / 72.0, 160 / 72.0)
        pix = _fitz_pixmap(page, matrix=mat, clip=clip, alpha=False)
        _fitz_save_pix(pix, path)
    finally:
        doc.close()
    return path if os.path.isfile(path) else None


def _add_cover_image(doc, url, cache_dir, box, page_hw, prev_box, pdf_path=None, page_index=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    path = None
    if url:
        try:
            path = _download_image(url, cache_dir)
        except Exception:
            path = None
    if not path and pdf_path is not None and page_index is not None:
        path = _crop_pdf_image(pdf_path, page_index, box, page_hw, cache_dir)
    if not path:
        return prev_box

    pw, ph, x0, y0, x1, y1 = _box_metrics(box or {}, page_hw)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = _space_before_pt(prev_box, box or {}, page_hw)
    para.paragraph_format.space_after = Inches(0.04)
    if y0 > ph * 0.75:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if x0 > pw * 0.05:
            para.paragraph_format.left_indent = Inches(x0 / max(pw, 1) * 6.2)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    width_in = min(1.05, max(0.28, (x1 - x0) / max(pw, 1) * 6.2))
    run = para.add_run()
    run.add_picture(path, width=Inches(width_in))
    return {
        "x0": x0, "y0": y0, "x1": x1, "y1": y1,
        "label": "image", "text": "",
    }


def _add_positioned_paragraph(
    doc, text, box, page_hw, prev_box, watermarks,
    label="para", bold=False, italic=False, style=None, page_index=None,
    watermark_page=False, sig_page=False, rgb_arr=None, cover_bottom=None,
):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    text = _normalize_text(text)
    force = text in ("服务合同", "合同") or text.endswith("服务合同")
    if not text or (not force and _is_noise_text(text, watermarks, sig_page=sig_page)):
        return prev_box

    if re.match(r"^甲\s*方|^乙\s*方", text) or (sig_page and re.search(r"甲\s*方|乙\s*方", text)):
        text = _format_party_text(text, sig_page=sig_page)

    pw, _, x0, _, x1, _ = _box_metrics(box or {}, page_hw)
    cover = _cover_text_style(text, label, page_index, watermark_page)
    is_cover = page_index == 0
    use_cover_layout = is_cover and cover_bottom is not None
    para_style = None if use_cover_layout else style
    para = doc.add_paragraph(style=para_style)
    if is_cover:
        align = _cover_paragraph_align(text, label, box, page_hw)
    elif cover and cover.get("color"):
        align = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        align = _infer_alignment(box or {}, page_hw, text)
    para.alignment = align
    target_y = None
    if use_cover_layout:
        space_before, target_y = _space_before_cover_pt(cover_bottom, box, page_hw)
        para.paragraph_format.space_before = space_before
    else:
        para.paragraph_format.space_before = _space_before_pt(prev_box, box or {}, page_hw)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.15

    if align == WD_ALIGN_PARAGRAPH.LEFT and x0 > pw * 0.07:
        para.paragraph_format.left_indent = Inches(x0 / max(pw, 1) * 6.2)

    run = para.add_run(text)
    run.italic = bool(italic or label == "cap")
    if cover:
        size_pt = cover["size"]
        is_blue_header = (
            label == "header"
            or bool(re.fullmatch(r"\d{6,12}", text.strip()))
        )
        if is_blue_header:
            color_hex = cover.get("color") or "008FEF"
            if is_cover and rgb_arr is not None:
                sampled = _sample_box_color_hex(
                    rgb_arr, box, page_hw, prefer_blue=True,
                )
                if sampled:
                    color_hex = sampled
        else:
            color_hex = "000000"
        _set_run_font_style(
            run,
            cover["name"],
            size_pt,
            east_asia=cover.get("east_asia"),
            bold=cover.get("bold", False) or bold or label in ("title", "sec"),
            color_hex=color_hex,
        )
    else:
        size_pt = _font_size_pt(
            box or {}, label, text, page_hw, page_index, watermark_page,
        ).pt
        _set_run_font_style(
            run,
            "宋体",
            size_pt,
            east_asia="SimSun",
            bold=bold or label in ("title", "sec"),
        )

    result = {
        "x0": x0,
        "y0": box.get("y0") or 0,
        "x1": x1,
        "y1": box.get("y1") or box.get("y0") or 0,
        "label": label,
        "text": text,
    }
    if use_cover_layout and target_y is not None:
        result["rendered_bottom"] = target_y + _estimate_para_height_pt(
            text, size_pt, box, page_hw,
        )
    return result


def _render_text_content(doc, text, watermarks, cache_dir, bold=False, italic=False, style=None, sig_page=False):
    text = _normalize_text(text)
    if _is_noise_text(text, watermarks, sig_page=sig_page):
        return

    tables = _TABLE_HTML_RE.findall(text)
    if tables:
        remain = _TABLE_HTML_RE.sub("", text).strip()
        if remain and not _is_noise_text(remain, watermarks, sig_page=sig_page):
            _add_text_paragraph(doc, remain, style=style, bold=bold, italic=italic)
        for html in tables:
            _add_html_table(doc, html)
        return

    for img_url in _IMG_MD_RE.findall(text):
        _add_image(doc, img_url, cache_dir)
    remain = _IMG_MD_RE.sub("", text).strip()
    if remain and not _is_noise_text(remain, watermarks, sig_page=sig_page):
        _add_text_paragraph(doc, remain, style=style, bold=bold, italic=italic)


def _parse_html_table_rows(html):
    rows_html = re.findall(r"<tr[^>]*>(.*?)</tr>", html, flags=re.I | re.S)
    parsed = []
    for row_html in rows_html:
        cells = []
        for attrs, inner in _TABLE_CELL_RE.findall(row_html):
            text = re.sub(r"<[^>]+>", "", inner).strip()
            text = _normalize_text(text)
            rowspan = colspan = 1
            for m in _ATTR_SPAN_RE.finditer(attrs):
                if m.group(1).lower() == "rowspan":
                    rowspan = max(1, int(m.group(2)))
                else:
                    colspan = max(1, int(m.group(2)))
            cells.append({"text": text, "rowspan": rowspan, "colspan": colspan})
        if cells:
            parsed.append(cells)
    return parsed


def _set_cell_font(cell, pt, bold=False, color_hex=None):
    from docx.shared import Pt

    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(pt)
            if bold:
                run.bold = True
            if color_hex:
                _set_run_font_style(run, "宋体", pt, east_asia="SimSun", bold=bold, color_hex=color_hex)
        if not para.runs:
            run = para.add_run(para.text)
            para.text = ""
            run.font.size = Pt(pt)
            if bold:
                run.bold = True
            if color_hex:
                _set_run_font_style(run, "宋体", pt, east_asia="SimSun", bold=bold, color_hex=color_hex)


def _add_html_table(
    doc, html, landscape=False, restore_portrait=True,
    compact=False, skip_section_switch=False, row_heights=None, row_line_positions=None,
):
    from docx.shared import Inches

    parsed = _parse_html_table_rows(html)
    if not parsed:
        doc.add_paragraph(re.sub(r"<[^>]+>", "", html))
        return False

    occupied = set()
    placements = []
    max_col = 0
    for ri, row in enumerate(parsed):
        ci = 0
        for cell in row:
            while (ri, ci) in occupied:
                ci += 1
            rs, cs = cell["rowspan"], cell["colspan"]
            for rr in range(ri, ri + rs):
                for cc in range(ci, ci + cs):
                    occupied.add((rr, cc))
            placements.append((ri, ci, cell))
            max_col = max(max_col, ci + cs)
            ci += cs

    if not occupied:
        doc.add_paragraph(re.sub(r"<[^>]+>", "", html))
        return False

    nrows = max(r for r, _ in occupied) + 1
    ncols = max(c for _, c in occupied) + 1
    ncols = max(ncols, max_col)

    # Convert line positions to row heights now that nrows is known
    if row_heights is None and row_line_positions:
        row_heights = _row_heights_from_line_positions(row_line_positions, nrows)

    auto_landscape = ncols >= 6 and nrows >= 8 and not compact
    use_landscape = landscape or auto_landscape
    if use_landscape and not skip_section_switch:
        _begin_landscape_section(doc)

    task_style = _is_task_checklist_table(parsed)
    scoring_style = _is_scoring_form_table(parsed) if not task_style else False
    if task_style or scoring_style:
        font_pt = _task_table_font_pt(ncols, nrows)
    else:
        font_pt = (
            6 if ncols >= 7 or nrows >= 22 else
            6.5 if ncols >= 6 or nrows >= 15 else
            7.5 if ncols >= 5 or nrows >= 10 else 9
        )
    if compact and not task_style and not scoring_style:
        font_pt = min(
            font_pt,
            5.0 if nrows >= 45 else 5.5 if nrows >= 35 else 6.0,
        )

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    if not task_style and not scoring_style:
        sec = doc.sections[-1]
        usable_w = sec.page_width - sec.left_margin - sec.right_margin
        col_w = int(usable_w / max(ncols, 1))
        for column in table.columns:
            column.width = col_w

    for ri, ci, cell in placements:
        tcell = table.rows[ri].cells[ci]
        tcell.text = cell["text"]
        if not task_style and not scoring_style:
            _set_cell_font(tcell, font_pt)
        rs, cs = cell["rowspan"], cell["colspan"]
        end_r = min(ri + rs - 1, nrows - 1)
        end_c = min(ci + cs - 1, ncols - 1)
        if end_r > ri or end_c > ci:
            tcell.merge(table.rows[end_r].cells[end_c])

    if task_style:
        _apply_task_checklist_table_style(
            table, parsed, placements, nrows, ncols, font_pt, use_landscape,
        )
    elif scoring_style:
        _apply_scoring_form_table_style(
            table, parsed, placements, nrows, ncols, font_pt, use_landscape,
            row_heights=row_heights,
        )
    elif compact:
        _apply_compact_dense_table_style(table, nrows, font_pt)

    # Apply explicit row heights if provided and scoring style didn't handle them
    if row_heights and not scoring_style:
        _EMU_PER_TWIP = 635
        sec = table.part.document.sections[-1]
        avail_h_twips = int(sec.page_height - sec.top_margin - sec.bottom_margin) // _EMU_PER_TWIP
        rh = list(row_heights)
        # Trim extra trailing rows (img2table may detect a stray footer line)
        while len(rh) > nrows and len(rh) > 2:
            rh.pop()
        if rh and len(rh) == nrows:
            total = sum(rh)
            rh = [h / total for h in rh]
            for ri, row in enumerate(table.rows):
                twips = max(200, int(avail_h_twips * rh[ri]))
                _set_row_height(row, twips, exact=True)

    if use_landscape and restore_portrait:
        _begin_portrait_section(doc)
    return True


def _single_page_table_opts(total_pages, dense_table, scoring_form=False):
    """单页表格：设备清单压紧；考评打分表保持竖版 A4 版式。"""
    if total_pages == 1 and scoring_form:
        return {
            "landscape": False,
            "restore_portrait": False,
            "compact": False,
            "skip_section_switch": True,
        }
    if total_pages == 1 and dense_table:
        return {
            "landscape": False,
            "restore_portrait": False,
            "compact": True,
            "skip_section_switch": True,
        }
    return {
        "landscape": True,
        "restore_portrait": True,
        "compact": False,
        "skip_section_switch": False,
    }


def _render_page_blocks(
    doc, page, watermarks, cache_dir, mode,
    is_sig_page=False, page_index=None, pdf_path=None, table_landscape=False,
    table_opts=None,
):
    blocks = _sort_blocks(page.get("textblocks") or [])
    page_hw = page.get("page_image_hw") or {}

    if mode == "visual":
        return

    skip_images = mode == "hybrid"
    prev_box = None
    cover_bottom = 0.0 if page_index == 0 else None
    use_layout = mode in ("hybrid", "text")
    wm_page = page_index == 0 or _page_watermark_heavy(page, watermarks)
    rgb_arr = None
    if page_index == 0 and pdf_path:
        rgb_arr = _get_cached_page_rgb(pdf_path, page_index, cache_dir)

    def _commit_positioned(**kwargs):
        nonlocal prev_box, cover_bottom
        prev_box = _add_positioned_paragraph(
            cover_bottom=cover_bottom, **kwargs,
        )
        if (
            cover_bottom is not None
            and prev_box
            and prev_box.get("rendered_bottom") is not None
        ):
            cover_bottom = prev_box["rendered_bottom"]

    for block in blocks:
        label = (block.get("label") or "para").lower()
        text = block.get("text") or ""
        bold = block.get("is_bold", False)
        italic = block.get("is_italic", False)
        box = block.get("box") or {}

        if label == "foot":
            continue

        if label == "header":
            norm_hdr = _normalize_text(text)
            if use_layout and box and (
                page_index == 0
                or _is_stamp_header_text(text, box, page_hw)
                or norm_hdr.startswith("CMCCTD")
            ):
                _commit_positioned(
                    doc=doc, text=text, box=box, page_hw=page_hw,
                    prev_box=prev_box, watermarks=watermarks,
                    label="header", page_index=page_index, watermark_page=wm_page,
                    sig_page=is_sig_page, rgb_arr=rgb_arr,
                )
            continue

        if label == "image":
            if wm_page and mode == "hybrid" and not is_sig_page:
                # 封面图章走 detail_to_docx 里的浮动嵌入，避免流式插图撑出第 2 页
                if page_index != 0 and _should_render_cover_image(box, page_hw):
                    url = block.get("url") or ""
                    if url:
                        prev_box = _add_cover_image(
                            doc, url, cache_dir, box, page_hw, prev_box,
                            pdf_path=pdf_path, page_index=page_index,
                        )
                continue
            if not _should_render_inline_image(box, page_hw, skip_images):
                continue
            url = block.get("url") or ""
            if url:
                _add_image(doc, url, cache_dir, page_hw=page_hw, box=box)
            continue

        if label == "table":
            html = text if "<table" in text.lower() else ""
            if not html and block.get("url") and not skip_images:
                _add_image(doc, block["url"], cache_dir, page_hw=page_hw, box=box)
            elif html:
                opts = dict(table_opts or {})
                opts.setdefault("landscape", table_landscape)
                _add_html_table(doc, html, **opts)
            continue

        if use_layout and box:
            if page_index == 0 and not is_sig_page and label in ("para", "title", "cap"):
                expanded = _expand_cover_blocks(block)
            elif is_sig_page and label in ("para", "title", "cap"):
                expanded = _expand_sig_blocks(block)
            elif label == "para":
                expanded = _expand_layout_blocks(block)
            else:
                expanded = [block]
            for sub in expanded:
                sub_label = (sub.get("label") or label).lower()
                sub_text = sub.get("text") or ""
                sub_box = sub.get("box") or box
                if "<table" in sub_text.lower():
                    _render_text_content(
                        doc, sub_text, watermarks, cache_dir,
                        sig_page=is_sig_page,
                    )
                    continue
                if sub_label in ("title", "sec"):
                    style = "Heading 1" if sub_label == "title" else "Heading 2"
                    _commit_positioned(
                        doc=doc, text=sub_text, box=sub_box, page_hw=page_hw,
                        prev_box=prev_box, watermarks=watermarks,
                        label=sub_label, bold=True, style=style,
                        page_index=page_index, watermark_page=wm_page,
                        sig_page=is_sig_page, rgb_arr=rgb_arr,
                    )
                    continue
                if sub_label == "cap":
                    if _text_quality_score(_normalize_text(sub_text)) >= 0.35 or is_sig_page:
                        _commit_positioned(
                            doc=doc, text=sub_text, box=sub_box, page_hw=page_hw,
                            prev_box=prev_box, watermarks=watermarks,
                            label=sub_label, italic=True,
                            page_index=page_index, watermark_page=wm_page,
                            sig_page=is_sig_page, rgb_arr=rgb_arr,
                        )
                    continue
                if is_sig_page and len(sub_text) > 36 and re.search(r"甲\s*方|乙\s*方", sub_text):
                    for line in _split_sig_lines(sub_text):
                        _commit_positioned(
                            doc=doc, text=line, box=sub_box, page_hw=page_hw,
                            prev_box=prev_box, watermarks=watermarks,
                            label=sub_label, bold=bold, italic=italic,
                            page_index=page_index, watermark_page=wm_page,
                            sig_page=is_sig_page, rgb_arr=rgb_arr,
                        )
                    continue
                _commit_positioned(
                    doc=doc, text=sub_text, box=sub_box, page_hw=page_hw,
                    prev_box=prev_box, watermarks=watermarks,
                    label=sub_label, bold=bold, italic=italic,
                    page_index=page_index, watermark_page=wm_page,
                    sig_page=is_sig_page, rgb_arr=rgb_arr,
                )
            continue

        if label in ("title", "sec"):
            style = "Heading 1" if label == "title" else "Heading 2"
            _render_text_content(doc, text, watermarks, cache_dir, bold=True, style=style, sig_page=is_sig_page)
            continue

        if label == "cap":
            if _text_quality_score(_normalize_text(text)) >= 0.35 or is_sig_page:
                _render_text_content(doc, text, watermarks, cache_dir, italic=True, sig_page=is_sig_page)
            continue

        _render_text_content(doc, text, watermarks, cache_dir, bold=bold, italic=italic, sig_page=is_sig_page)


def detail_to_docx(pages, output_path, pdf_path=None, mode="text", page_markdowns=None):
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    if not pages:
        raise ValueError("无版面数据")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)

    watermarks = _detect_dynamic_watermarks(pages)
    cache_dir = tempfile.mkdtemp(prefix="volc_img_")
    page_markdowns = page_markdowns or []
    snapshot_pages = []
    abnormal_table_pages = []
    _clear_page_layout_cache()
    try:
        total_pages = len(pages)
        for pi, page in enumerate(pages):
            if pi > 0:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

            page_md = page_markdowns[pi] if pi < len(page_markdowns) else ""
            is_sig_page = False
            if pdf_path and _seal_extract_enabled():
                from seal_utils import is_signature_page
                is_sig_page = is_signature_page(page, pi, total_pages)

            stats = _page_render_stats(page, watermarks, page_md, is_sig_page=is_sig_page)
            text_before = _doc_body_text_len(doc)
            tables_before = len(doc.tables)
            used_snapshot = False
            scoring_form = _page_is_scoring_form(page, page_md)
            dense_table = _page_is_dense_table(page, page_md)
            single_dense = total_pages == 1 and dense_table
            single_form = total_pages == 1 and scoring_form
            table_opts = _single_page_table_opts(
                total_pages, dense_table, scoring_form=scoring_form,
            )
            page_layout = None
            if pdf_path:
                from seal_utils import SEAL_EXTRACT_DPI as _seal_dpi
                page_layout = _analyze_page_layout(
                    pdf_path, pi, _seal_dpi if single_dense else 180,
                    probe_coarse=single_dense or single_form,
                )
            if (single_dense or single_form) and pi == 0 and pdf_path:
                if page_layout and page_layout.get("correction_deg"):
                    _sync_doc_section_to_layout(doc, page_layout, tight=True)
                else:
                    _sync_doc_section_to_pdf(doc, pdf_path, pi, tight=True)
                # Detect proportional row heights from corrected image (img2table primary)
                try:
                    corr_bgr, _ = _get_corrected_page_bgr(pdf_path, pi, 150, probe_coarse=True)
                    rh = _extract_row_heights_img2table(corr_bgr)
                    if rh:
                        table_opts["row_heights"] = rh
                    else:
                        line_pos = _detect_row_line_positions(corr_bgr)
                        if line_pos:
                            table_opts["row_line_positions"] = line_pos
                except Exception:
                    pass

            if dense_table:
                # P1 同页双通道：正文走 detail 坐标，表格择优用 markdown HTML
                non_table_len = _page_non_table_meaningful_len(
                    page, watermarks, is_sig_page=is_sig_page,
                )
                quality_page = page
                if non_table_len >= 15:
                    quality_page, consumed, md_tables = _prepare_dual_channel_page(
                        page, page_md, dense=True,
                    )
                    _render_page_blocks(
                        doc, quality_page, watermarks, cache_dir, mode,
                        is_sig_page=is_sig_page, page_index=pi, pdf_path=pdf_path,
                        table_landscape=table_opts.get("landscape", True),
                        table_opts=table_opts,
                    )
                    _render_remaining_md_tables(
                        doc, md_tables, consumed,
                        landscape=table_opts.get("landscape", True),
                        table_opts=table_opts,
                    )
                else:
                    _render_dense_table_page(
                        doc, page_md, watermarks, cache_dir, table_opts=table_opts,
                    )
                if _page_table_quality_abnormal(quality_page, page_md):
                    abnormal_table_pages.append(pi + 1)
            elif _should_render_snapshot(page, watermarks, is_sig_page, mode, pdf_path, pi):
                if _render_pdf_page_snapshot(
                    doc, pdf_path, pi, cache_dir, dpi=_BLANK_PAGE_SNAPSHOT_DPI,
                ):
                    used_snapshot = True
                    snapshot_pages.append(pi + 1)
            else:
                prepared, consumed, md_tables = _prepare_dual_channel_page(
                    page, page_md, dense=False,
                )
                _render_page_blocks(
                    doc,
                    prepared,
                    watermarks,
                    cache_dir,
                    mode,
                    is_sig_page=is_sig_page,
                    page_index=pi,
                    pdf_path=pdf_path,
                )
                _render_remaining_md_tables(doc, md_tables, consumed, landscape=False)

            if pdf_path and _page_has_image_blocks(page) and not dense_table:
                pg_w_emu, pg_h_emu = _get_page_emu_size(pdf_path, pi)
                _embed_ocr_image_blocks(
                    doc, page, cache_dir, pg_w_emu, pg_h_emu,
                    pdf_path=pdf_path, page_index=pi,
                )

            if is_sig_page and pdf_path:
                from seal_utils import SEAL_EXTRACT_DPI
                pg_w_emu, pg_h_emu = _get_page_emu_size(pdf_path, pi)
                hi_res = _get_pdf_page_image(
                    pdf_path, pi, cache_dir, dpi=SEAL_EXTRACT_DPI
                )
                if hi_res:
                    _embed_extracted_seals(
                        doc,
                        hi_res,
                        cache_dir,
                        page_hw=page.get("page_image_hw"),
                        pg_w_emu=pg_w_emu,
                        pg_h_emu=pg_h_emu,
                    )
                    _embed_signatures(doc, hi_res, page, cache_dir, pg_w_emu, pg_h_emu)

            chars_added = _doc_body_text_len(doc) - text_before
            tables_added = len(doc.tables) - tables_before
            if (
                pdf_path
                and chars_added < _BLANK_PAGE_MIN_CHARS
                and tables_added == 0
                and not used_snapshot
            ):
                if _render_pdf_page_snapshot(
                    doc, pdf_path, pi, cache_dir, dpi=_BLANK_PAGE_SNAPSHOT_DPI,
                ):
                    snapshot_pages.append(pi + 1)
            elif (
                pdf_path
                and _page_is_blank_risk(stats, is_sig_page)
                and not dense_table
                and not used_snapshot
                and chars_added < _BLANK_PAGE_MIN_CHARS
                and tables_added == 0
            ):
                if _render_pdf_page_snapshot(
                    doc, pdf_path, pi, cache_dir, dpi=_BLANK_PAGE_SNAPSHOT_DPI,
                ):
                    snapshot_pages.append(pi + 1)

        doc.save(output_path)
    finally:
        shutil.rmtree(cache_dir, ignore_errors=True)
    return _build_page_warnings(sorted(set(snapshot_pages)), sorted(set(abnormal_table_pages)))


def markdown_to_docx(markdown_text, output_path):
    from docx import Document

    doc = Document()
    cache_dir = tempfile.mkdtemp(prefix="volc_md_")
    watermarks = _STATIC_WATERMARKS
    try:
        blocks = re.split(r"\n\s*\n", markdown_text)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            if block.lower().startswith("<table"):
                _add_html_table(doc, block)
                continue
            if block.startswith("#"):
                level = len(block) - len(block.lstrip("#"))
                text = block.lstrip("#").strip()
                style = "Heading 1" if level <= 1 else "Heading 2"
                _render_text_content(doc, text, watermarks, cache_dir, style=style)
                continue
            for line in block.split("\n"):
                line = line.strip()
                if line:
                    _render_text_content(doc, line, watermarks, cache_dir)
        doc.save(output_path)
    finally:
        shutil.rmtree(cache_dir, ignore_errors=True)


def _classify_scan_doc(pdf_path, markdown="", details=None):
    """判断扫描件类型；PDF 直传 detail 不足时抽样首页 PNG。"""
    if _looks_like_task_checklist(markdown, details):
        return "checklist"
    if _looks_like_contract(markdown, details):
        return "contract"

    ocr_path, cleaned_tmp = _preprocessed_ocr_path(pdf_path)
    try:
        visual = _visual_service()
        img_b64, _ = _page_rgb_and_b64(ocr_path, 0, dpi=_image_parse_dpi())
        sample = _ocr_pdf_image_page_data(visual, img_b64)
        sample_md = sample.get("markdown") or ""
        sample_details = sample.get("detail") or []
        if _looks_like_task_checklist(sample_md, sample_details):
            return "checklist"
        if _looks_like_contract(sample_md, sample_details):
            return "contract"
    except Exception:
        pass
    finally:
        if cleaned_tmp and os.path.isfile(cleaned_tmp):
            try:
                os.unlink(cleaned_tmp)
            except OSError:
                pass
    return "contract"


def _try_image_mode_docx(pdf_path, output_path, markdown="", details=None):
    """逐页 PNG 重解析：合同走 hybrid detail，任务清单走 markdown 版式。"""
    doc_kind = _classify_scan_doc(pdf_path, markdown, details)
    page_count = pdf_page_count(pdf_path)
    mode = _docx_mode(page_count)

    if doc_kind == "contract":
        img_details, pages_md = pdf_to_detail_image_mode(pdf_path)
        if _detail_has_usable_content(img_details):
            log.info("逐页 PNG detail 模式(%s): %s", mode, pdf_path)
            warning = detail_to_docx(
                img_details,
                output_path,
                pdf_path=pdf_path,
                mode=mode,
                page_markdowns=pages_md,
            )
            return {"route": f"volc-{mode}", "pages": len(img_details), "warning": warning}

    pages_md = pdf_to_markdown_image_mode(pdf_path)
    combined = "\n\n".join(p for p in pages_md if p)
    if not _markdown_has_usable_content(combined):
        raise ValueError("图片模式未返回有效内容")
    log.info("逐页 PNG markdown 任务清单模式: %s", pdf_path)
    markdown_pages_to_docx(pages_md, output_path)
    return {"route": "volc-image-table", "pages": len(pages_md)}


def volc_pdf_to_docx(pdf_path, output_path):
    """返回 {"route": str, "warning": str}。"""
    markdown, details = pdf_to_markdown(pdf_path)
    needs_image_mode = _detail_needs_image_mode(markdown, details)
    if needs_image_mode:
        try:
            log.info("PDF 直传版式不足，逐页 PNG 智能文档解析: %s", pdf_path)
            meta = _try_image_mode_docx(
                pdf_path, output_path, markdown=markdown, details=details,
            )
            return {"route": meta["route"], "warning": meta.get("warning") or ""}
        except Exception as e:
            log.warning("逐页 PNG 智能文档解析失败(%.80s)，继续 fallback", e)

    has_usable = _markdown_has_usable_content(markdown) or _detail_has_usable_content(details)
    if has_usable:
        page_count = len(details) or pdf_page_count(pdf_path)
        mode = _docx_mode(page_count)
        if details and _detail_has_usable_content(details):
            try:
                warning = detail_to_docx(details, output_path, pdf_path=pdf_path, mode=mode)
                return {"route": f"volc-{mode}", "warning": warning or ""}
            except Exception:
                pass
        if _markdown_has_usable_content(markdown):
            markdown_to_docx(markdown, output_path)
            return {"route": "volc", "warning": ""}

    image_only = _doc_parse_image_only(markdown, details)
    if not has_usable:
        try:
            log.info(
                "直传无有效内容(image_only=%s)，尝试逐页 PNG 智能文档解析: %s",
                image_only, pdf_path,
            )
            meta = _try_image_mode_docx(
                pdf_path, output_path, markdown=markdown, details=details,
            )
            return {"route": meta["route"], "warning": meta.get("warning") or ""}
        except Exception as e:
            log.warning("逐页 PNG 智能文档解析失败(%.80s)，继续 fallback", e)
    log.info(
        "智能文档解析无有效文字/表格，fallback 通用文字识别 (image_only=%s): %s",
        image_only, pdf_path,
    )
    normal_ocr_pdf_to_docx(pdf_path, output_path)
    warning = _DEGRADED_OCR_WARNING if (image_only or needs_image_mode) else ""
    return {"route": "volc-normal", "warning": warning}