#!/usr/bin/env python3
"""C.pdf -> Word via Volc 通用文字识别 (ocr_normal)."""
import os
import sys
import time

sys.path.insert(0, os.path.dirname(__file__))
from ssh_helper import run, put, fetch

PDF_LOCAL = r"C:\Users\paz\Desktop\P T W 测试\C.pdf"
OUT_LOCAL = r"C:\Users\paz\Desktop\P T W 测试\C_normal.docx"
REMOTE_PDF = "/tmp/C.pdf"
REMOTE_DOCX = "/tmp/C_normal.docx"
DPI = 200

print("=== Upload C.pdf ===")
put(PDF_LOCAL, REMOTE_PDF)

print("\n=== ocr_normal -> Word ===")
t0 = time.time()
cmd = f"""
cd /home/toolbox && TOOLBOX_ENV=/home/toolbox/toolbox.env python3.8 - <<'PYEOF'
import os, sys, json, base64
sys.path.insert(0, "/home/toolbox/backend")
os.environ["TOOLBOX_ENV"] = "/home/toolbox/toolbox.env"

import fitz
import numpy as np
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor
from volc_ocr import volc_credentials
from volcengine.visual.VisualService import VisualService

PDF = "{REMOTE_PDF}"
OUT = "{REMOTE_DOCX}"
DPI = {DPI}

visual = VisualService()
ak, sk = volc_credentials()
visual.set_ak(ak)
visual.set_sk(sk)

doc_pdf = fitz.open(PDF)
page = doc_pdf[0]
pix = page.get_pixmap(matrix=fitz.Matrix(DPI / 72.0, DPI / 72.0), alpha=False)
png_path = "/tmp/C_page_normal.png"
pix.save(png_path)
doc_pdf.close()

img = Image.open(png_path).convert("RGB")
arr = np.asarray(img)

with open(png_path, "rb") as f:
    img_b64 = base64.b64encode(f.read()).decode("ascii")

resp = visual.ocr_normal({{"image_base64": img_b64}})
if not resp or resp.get("code") != 10000:
    raise SystemExit((resp or {{}}).get("message") or "ocr_normal failed")

data = resp.get("data") or {{}}
lines = data.get("line_texts") or []
chars = data.get("chars") or []
print("lines", len(lines), "chars_groups", len(chars))

def px_to_pt(h):
    return max(9, min(28, round(h * 72.0 / DPI * 0.95, 1)))

def sample_blue(line_chars):
    blues = []
    for c in line_chars or []:
        x = int(c.get("x", 0) + c.get("width", 0) / 2)
        y = int(c.get("y", 0) + c.get("height", 0) / 2)
        if 0 <= y < arr.shape[0] and 0 <= x < arr.shape[1]:
            r, g, b = arr[y, x]
            if b > max(r, g) + 35 and b > 120:
                blues.append((r, g, b))
    if not blues:
        return None
    r = int(sum(t[0] for t in blues) / len(blues))
    g = int(sum(t[1] for t in blues) / len(blues))
    b = int(sum(t[2] for t in blues) / len(blues))
    return RGBColor(r, g, b)

def avg_h(line_chars):
    if not line_chars:
        return 0
    return sum(c.get("height", 0) for c in line_chars) / len(line_chars)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal.font.size = Pt(10.5)

for i, text in enumerate(lines):
    text = (text or "").strip()
    if not text:
        continue
    line_chars = chars[i] if i < len(chars) else []
    h = avg_h(line_chars)
    is_title = i == 0 or h >= 55
    is_last_blue = "故乡" in text and "梦里" in text
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "宋体"
    try:
        from docx.oxml.ns import qn
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        pass
    if is_title:
        run.bold = True
        run.font.size = Pt(px_to_pt(h) if h else 22)
        color = sample_blue(line_chars) or RGBColor(0, 114, 239)
        run.font.color.rgb = color
    elif is_last_blue:
        run.font.size = Pt(px_to_pt(h) if h else 11)
        color = sample_blue(line_chars) or RGBColor(0, 114, 239)
        run.font.color.rgb = color
    else:
        run.font.size = Pt(px_to_pt(h) if h else 11)

doc.save(OUT)
print("saved", OUT, os.path.getsize(OUT))
print("preview", "\\n".join(lines[:3]))
PYEOF
"""
code, out, err = run(cmd, timeout=120)
elapsed = int(time.time() - t0)
print(f"exit={code} wall={elapsed}s")
print(out)
if err.strip():
    print("stderr:", err.strip()[:500])
if code != 0:
    sys.exit(1)

print("\n=== Download ===")
fetch(REMOTE_DOCX, OUT_LOCAL)
print("saved:", OUT_LOCAL)
print("size_kb:", round(os.path.getsize(OUT_LOCAL) / 1024, 1))