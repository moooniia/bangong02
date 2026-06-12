#!/usr/bin/env python3
"""部署前冒烟：语法检查 + 五样例回归（page_1 / page_7 / A / B / C）。"""
import os
import subprocess
import sys
import tempfile
import zipfile

ROOT = os.path.dirname(os.path.abspath(__file__))
BACKEND = os.path.join(ROOT, "server", "backend")
SAMPLE_DIR = os.environ.get(
    "SMOKE_SAMPLE_DIR",
    r"C:\Users\paz\Desktop\P T W 测试",
)
API_BASE = os.environ.get("SMOKE_API_BASE", "http://139.196.28.78")

PY_FILES = [
    "volc_ocr.py",
    "app.py",
    "preprocessing.py",
    "seal_utils.py",
]

# 够用标准：路由正确 + 有可编辑内容 + 不明显退化（非像素级）
REGRESSION_CASES = [
    {
        "name": "page_1.pdf",
        "filename": "page_1.pdf",
        "expect_route": "volc-hybrid",
        "timeout": 180,
        "checks": {
            "min_text_paragraphs": 5,
            "max_page_breaks": 0,
            "min_chars": 80,
            "min_anchors": 1,
            "must_contain": ("服务合同", "甲方", "乙方"),
        },
    },
    {
        "name": "A.pdf",
        "filename": "A.pdf",
        "expect_route": "volc-hybrid",
        "timeout": 600,
        "checks": {
            "min_chars": 800,
            "min_text_paragraphs": 20,
            "must_contain": ("甲方", "乙方", "服务合同"),
        },
    },
    {
        "name": "B.pdf",
        "filename": "B.pdf",
        "expect_route": "volc-image-table",
        "timeout": 300,
        "checks": {
            "min_chars": 100,
            "min_tables": 1,
            "must_contain": ("任务",),
        },
    },
    {
        "name": "C.pdf",
        "filename": "C.pdf",
        "expect_route": "volc-normal",
        "timeout": 180,
        "checks": {
            "min_chars": 50,
            "min_text_paragraphs": 3,
        },
    },
    {
        "name": "page_7.pdf",
        "filename": "page_7.pdf",
        "expect_route": "volc-hybrid",
        "timeout": 180,
        "checks": {
            "min_chars": 150,
            "min_text_paragraphs": 6,
            "max_page_breaks": 0,
            "must_contain": ("第八条", "第九条", "甲方", "乙方"),
        },
    },
]


def compile_check() -> tuple[bool, str]:
    errors = []
    for name in PY_FILES:
        path = os.path.join(BACKEND, name)
        if not os.path.isfile(path):
            errors.append(f"缺少 {name}")
            continue
        r = subprocess.run(
            [sys.executable, "-m", "py_compile", path],
            capture_output=True,
            text=True,
        )
        if r.returncode != 0:
            errors.append(f"{name}: {r.stderr.strip() or r.stdout.strip()}")
    if errors:
        return False, "编译检查失败:\n" + "\n".join(errors)
    return True, f"编译检查通过 ({len(PY_FILES)} 个核心文件)"


def inspect_docx(path: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("回归需要 python-docx，请 pip install python-docx")

    doc = Document(path)
    text_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    body_text = "".join(p.text for p in doc.paragraphs)
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    return {
        "text_paragraphs": len(text_paragraphs),
        "page_breaks": xml.count('w:type="page"'),
        "anchors": xml.count("wp:anchor"),
        "tables": len(doc.tables),
        "chars": len(body_text.strip()),
        "body_text": body_text,
    }


def _check_metrics(case: dict, metrics: dict) -> list[str]:
    checks = case.get("checks") or {}
    errors = []
    if "expect_route" in case:
        pass
    if metrics["chars"] < checks.get("min_chars", 0):
        errors.append(f"字数不足: {metrics['chars']} < {checks['min_chars']}")
    if metrics["text_paragraphs"] < checks.get("min_text_paragraphs", 0):
        errors.append(
            f"正文段过少: {metrics['text_paragraphs']} < {checks['min_text_paragraphs']}"
        )
    if metrics["tables"] < checks.get("min_tables", 0):
        errors.append(f"表格过少: {metrics['tables']} < {checks['min_tables']}")
    if metrics["page_breaks"] > checks.get("max_page_breaks", 9999):
        errors.append(
            f"分页过多: {metrics['page_breaks']} > {checks['max_page_breaks']}"
        )
    if metrics["anchors"] < checks.get("min_anchors", 0):
        errors.append(f"浮动图过少: {metrics['anchors']} < {checks['min_anchors']}")
    for kw in checks.get("must_contain") or ():
        if kw not in metrics["body_text"]:
            errors.append(f"缺少关键字: {kw}")
    return errors


def regression_check() -> tuple[bool, str]:
    try:
        import requests
    except ImportError:
        return True, "跳过五样例回归（未安装 requests）"

    lines = []
    missing = []
    for case in REGRESSION_CASES:
        path = os.path.join(SAMPLE_DIR, case["filename"])
        if not os.path.isfile(path):
            missing.append(case["name"])
            continue

        url = f"{API_BASE}/api/convert"
        with open(path, "rb") as f:
            r = requests.post(
                url,
                files={"file": (case["filename"], f, "application/pdf")},
                data={"format": "docx"},
                timeout=case.get("timeout", 300),
            )
        if r.status_code != 200:
            return False, f"{case['name']} HTTP {r.status_code}: {r.text[:200]}"
        data = r.json()
        if not data.get("success"):
            return False, f"{case['name']} 转换失败: {data}"

        route = data.get("route", "")
        expect = case.get("expect_route", "")
        if expect and route != expect:
            return False, (
                f"{case['name']} 路由异常: 期望 {expect} 实际 {route}"
            )

        out_name = data.get("filename") or ""
        if not out_name:
            return False, f"{case['name']} 响应无 filename"

        dl = requests.get(
            f"{API_BASE}/api/download/{out_name}",
            timeout=120,
        )
        if dl.status_code != 200:
            return False, f"{case['name']} 下载失败 HTTP {dl.status_code}"

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        try:
            tmp.write(dl.content)
            tmp.close()
            metrics = inspect_docx(tmp.name)
            errs = _check_metrics(case, metrics)
            if errs:
                return False, f"{case['name']} 指标未达标:\n  " + "\n  ".join(errs)
            warn = data.get("warning") or ""
            warn_note = f" warning={warn[:40]}..." if warn else ""
            lines.append(
                f"OK {case['name']} -> {route} "
                f"(段={metrics['text_paragraphs']} 字={metrics['chars']} "
                f"表={metrics['tables']} 分页符={metrics['page_breaks']}){warn_note}"
            )
        finally:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

    if missing:
        lines.append("跳过缺失样例: " + ", ".join(missing))
    if not lines:
        return True, "五样例回归: 无可用样例文件"
    return True, "五样例回归:\n" + "\n".join(lines)


def run_smoke(regression: bool = False) -> tuple[bool, str]:
    ok, msg = compile_check()
    if not ok:
        return ok, msg
    parts = [msg]
    if regression:
        ok2, msg2 = regression_check()
        parts.append(msg2)
        if not ok2:
            return False, "\n".join(parts)
    return True, "\n".join(parts)


# 兼容旧调用
def online_abc_check() -> tuple[bool, str]:
    return regression_check()


if __name__ == "__main__":
    import argparse

    p = argparse.ArgumentParser(description="部署前冒烟与四样例回归")
    p.add_argument(
        "--regression",
        action="store_true",
        help="跑 page_1 / A / B / C 四样例线上回归",
    )
    p.add_argument(
        "--online",
        action="store_true",
        help="同 --regression（兼容旧参数）",
    )
    args = p.parse_args()
    success, text = run_smoke(regression=args.regression or args.online)
    print(text)
    sys.exit(0 if success else 1)