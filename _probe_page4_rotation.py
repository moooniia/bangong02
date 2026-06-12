#!/usr/bin/env python3
"""Probe page_4.pdf rotation signals vs WPS baseline."""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "server", "backend"))

PDF = r"C:\Users\paz\Desktop\P T W 测试\page_4.pdf"
WPS = r"C:\Users\paz\Desktop\P T W 测试\page_4.docx"
OURS = r"C:\Users\paz\Desktop\P T W 测试\page_4_v0.10.12.docx"

import fitz

doc = fitz.open(PDF)
page = doc[0]
print("=== PDF meta ===")
print("rotation meta:", page.rotation)
print("mediabox:", page.rect)
print("pix w/h @150dpi:", end=" ")
pix = page.get_pixmap(matrix=fitz.Matrix(150 / 72, 150 / 72))
print(pix.width, pix.height)
doc.close()

import volc_ocr as vo

bgr = vo._render_page_bgr(PDF, 0, 150)
print("\n=== Current heuristics @150dpi ===")
print("ink cover:", vo._ink_cover_metrics(bgr))
print("sidebar vertical:", vo._has_sidebar_vertical_layout(bgr))
print("upright header score:", vo._upright_header_score(bgr))
print("visual sideways:", vo._detect_visual_sideways_rotation(bgr))
print("coarse allow_full=False:", vo._detect_best_coarse_rotation(bgr, allow_full_probe=False))
print("coarse allow_full=True:", vo._detect_best_coarse_rotation(bgr, allow_full_probe=True))
print("fine skew:", vo._detect_fine_skew_deg(bgr))
for deg in (0, 90, 180, 270):
    rot = vo._rotate_bgr(bgr, deg) if deg else bgr
    print(
        f"  deg {deg:3d}: align={vo._orientation_alignment_score(rot):.1f} "
        f"header={vo._upright_header_score(rot):.1f} "
        f"combined={vo._orientation_combined_score(rot):.1f} "
        f"sideways={vo._detect_visual_sideways_rotation(rot)}"
    )

layout_probe = vo._analyze_page_layout(PDF, 0, 150, probe_coarse=False)
layout_full = vo._analyze_page_layout(PDF, 0, 150, probe_coarse=True)
print("\nlayout probe_coarse=False:", layout_probe)
print("layout probe_coarse=True:", layout_full)

# Compare docx outputs
from docx import Document


def docx_stats(path, label):
    if not os.path.isfile(path):
        print(f"\n=== {label}: MISSING {path} ===")
        return
    d = Document(path)
    secs = d.sections
    print(f"\n=== {label} ===")
    print("sections:", len(secs))
    if secs:
        s = secs[0]
        print(
            "page: %.1fx%.1f in, landscape=%s"
            % (s.page_width.inches, s.page_height.inches, s.orientation)
        )
    chars = sum(len(p.text) for p in d.paragraphs)
    tbl_chars = sum(len(c.text) for t in d.tables for r in t.rows for c in r.cells)
    print("paragraph chars:", chars, "table chars:", tbl_chars, "tables:", len(d.tables))
    if d.tables:
        t = d.tables[0]
        print("table shape:", len(t.rows), "x", len(t.columns))
        sample = []
        for ri in range(min(3, len(t.rows))):
            row = [c.text.strip()[:20] for c in t.rows[ri].cells[:5]]
            sample.append(row)
        for row in sample:
            print(" ", row)


docx_stats(WPS, "WPS page_4.docx")
docx_stats(OURS, "Our v0.10.12")